import pandas as pd  # pandas 是数据分析中不可或缺的工具,用于对sql查询的表格进行处理
import re  # re 是 Python 内置的正则表达式模块，用于对大模型的输出进行文本匹配
from functools import wraps  # 装饰器
from flask import Flask, jsonify, Response, request, redirect, url_for  # 导入flask的快捷库
import flask
import os
from cache import MemoryCache  # 利用字典来作为临时的内存存储
from vanna.chromadb import ChromaDB_VectorStore  # 导入Vanna里设置好的chromdb文件
from vanna.ollama import Ollama  # 导入Vanna配置好的ollama交互文件
from flask import Flask, request, send_file
from io import BytesIO  # BytesIO 是一个内存中的字节流对象，在这里用于把文件存在内存中，然后供前端下载
from docx import Document  # 用于生成文档的库
import json
from typing import List, Tuple, Union  # 导入一些数据类型
import ollama
from helpers import pptx_helper, text_helper  # 从helpers文件夹中导入生成ppt的函数

app = Flask(__name__, static_url_path='')

# SETUP
cache = MemoryCache()

# from vanna.local import LocalContext_OpenAI
# vn = LocalContext_OpenAI()

# from vanna.remote import VannaDefault
# vn = VannaDefault(model=os.environ['VANNA_MODEL'], api_key=os.environ['VANNA_API_KEY'])

from typing import List
from chromadb.utils.embedding_functions import EmbeddingFunction

cilent = ollama.Client(host="192.168.30.130")


# 自定义函数：用于在Vanna类里面使用自己定义的嵌入模型
class OllamaEmbeddingFunction(EmbeddingFunction[List[str]]):
    def __init__(self, model_name="nomic-embed-text:latest"):
        self.model_name = model_name

    def __call__(self, input: List[str]) -> List[List[float]]:
        embeddings = []
        for text in input:
            # 使用 ollama.embeddings 获取嵌入向量
            response = cilent.embeddings(
                model=self.model_name,
                prompt=text
            )
            embedding = response['embedding']  # 获取嵌入向量
            embeddings.append(embedding)

        return embeddings


collections = {}  # 这个是用于装载定义的数据库背景知识，包括数据表的ddl，对字段的解释等内容。


# 这个是自定义一个叫Vanna的类，集成ChromaDB_VectorStore和Ollama这两个类，用于本地部署数据库和大模型
class MyVanna(ChromaDB_VectorStore, Ollama):
    def __init__(self, config=None):
        ChromaDB_VectorStore.__init__(self, config=config)
        Ollama.__init__(self, config=config)

    # 这个函数是自定义的，只需要执行一次就行，下次再启动项目的时候，把if len(text)>1:这后面的代码注释掉
    # 这个函数是将collections里面装载的内容存到本地chromadb数据库中，因为第一次要把内容向量化，然后再存进数据库，所以第一次运行比较久
    # 第二次运行就没必要在存数据库了，所以就把if len(text)>1:这后面的代码注释掉（这种方式不太方便，应该优化）
    def train_yu(
            self,
            text=[],
            name=""
    ):
        if name:
            collections[name] = self.chroma_client.get_or_create_collection(
                name=name,
            )
        # if len(text)>=1:
        #     for i, d in enumerate(text):
        #         response = cilent.embeddings(model="bge-m3:latest", prompt=d)
        #         embedding = response["embedding"]
        #         collections[name].add(
        #             ids=[str(i)],
        #             embeddings=[embedding],
        #             documents=[d]
        #         )

    # 这个函数是自定义函数，用于将用户提问的问题，和数据库中每张表的固定的字段内容进行匹配，然后筛选出哪张表是最适合用户提问的问题的。
    def get_similar_text_yu(self, question):
        best_score = 1000
        best_data = []
        best_name = ""
        ddl_list = []
        doucment_list = []
        response = cilent.embeddings(
            prompt=question,
            model="bge-m3:latest"
        )

        for i in collections.keys():
            if i in question:
                results = collections[i].query(
                    query_embeddings=[response["embedding"]],
                    n_results=20,
                )
                data = results['documents'][0]
                score = results['distances'][0]
                best_name = i
                best_data = data
                break
            else:
                results = collections[i].query(
                    query_embeddings=[response["embedding"]],
                    n_results=20,
                )
                data = results['documents'][0]
                score = results['distances'][0]
                print(score)
                if score[0] < best_score:
                    best_name = i
                    best_score = score[0]
                    best_data = data

        if best_name == 'twb_ygl0078':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("mat_ind_name_5_str", ",".join(best_data))

        # tmp=data_rag[best_name][2]

        doucment_list.append(data_rag[best_name][0] + "\n" + data_rag[best_name][1] + "\n" + tmp)

        return [], doucment_list, tmp

    # 这个函数是对Vanna里面的函数重写，修改了一些流程
    def generate_sql(self, question: str, allow_llm_to_see_data=False, doc_list=[], ddl_list=[], **kwargs) -> list:

        if self.config is not None:
            initial_prompt = self.config.get("initial_prompt", None)
        else:
            initial_prompt = None

        # question_sql_list = self.get_similar_question_sql(question, **kwargs)
        question_sql_list = []  # 这个变量暂时不用了，暂时不需要提供  问题-sql答案  这样的文本了

        # ddl_list = self.get_related_ddl(question, **kwargs)
        context_tmp = ""
        # ddl_list=[]
        if len(doc_list) == 0:  # doc_list，ddl_list，用来实现sql自纠错机制
            ddl_list, doc_list, context_tmp = self.get_similar_text_yu(question)  # doc_list是system的内容
        # 在这里加上提示词,规范输出
        template_input = f"""
        
---

### 规则说明
1. **判断是否需要查询数据库**：
   - 如果问题涉及数据查询或计算，请先一步步分析问题，再生成 SQL。
   - 如果问题无需查询数据库，请直接回答用户问题。
2. **SQL 生成规则**：
   - 表名、字段和条件必须严格按照 system 提供的固定取值范围。
   - 使用 `WITH` 结构组织查询，以提高可读性和扩展性。
   - 确保查询在 SQL 中给出具体的结果。
   - **无论查询还是计算，若表中有单位字段，必须在最终结果中包含单位字段**。
3. **单位字段的特殊规则**：
   - 如果其他字段使用了聚合函数（如 `SUM`、`AVG` 等），单位字段应加入 `GROUP BY` 子句。
   - 如果计算结果与单位字段无直接关联（如跨列计算），结果表中仍需保留单位字段，作为辅助信息。
4. **特殊注意**：
   - 避免跨行计算，所有计算仅在单行内完成。
   - 日期字段类型为 `timestamp`，查询中应考虑日期格式兼容。
   - **比率计算时不再乘以 100，直接输出比率值**。
   - 如果查询中的字段发生别名变化（例如 AS 用法），应确保后续 SELECT 使用正确的别名字段，而不是原字段名
---

### 用户提问
{question}

---

### 输出要求
根据用户提问和规则说明：
1. 判断问题是否需要查询数据库。
   - 如果需要，请一步步分析问题后生成 SQL。
   - 如果不需要，请直接回答问题。
2. **生成 SQL 的结构**：
   - 使用 `WITH` 子句进行逻辑拆分。
   - 确保查询结果输出具体的计算或筛选值。
   - **无论查询还是计算，结果中必须包含单位字段，且不得引发 SQL 语法错误**。
   - **确保 SQL 查询中的公司和日期字段始终按需要进行分组，并避免跨行计算或计算错误**。
   - 如果有聚合函数，单位字段必须加入 `GROUP BY` 子句。
   - 日期字段类型为 `timestamp`，查询中应考虑日期格式兼容。
3. **SQL 示例格式**：
   ```sql
   -- 请将生成的 SQL 放置在此格式中，确保语法正确
   ```
"""
        # question+="，请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL，如果不需要，请直接回答，如果需要生成sql，一步一步的分析问题，生成sql来解决问题给出答案，最好使用with的结构，请严格按照system内容的固定取值列表来生成sql，注意一定要在sql中给出结果。"
        prompt = self.get_sql_prompt(  # 将自己定义的问题，制作成可以直接输入到ollama里去，与大模型交流的格式
            initial_prompt=initial_prompt,
            question=template_input,
            question_sql_list=question_sql_list,
            ddl_list=ddl_list,
            doc_list=doc_list,
            **kwargs,
        )
        self.log(title="SQL Prompt", message=prompt)
        llm_response = self.submit_prompt(prompt, **kwargs)  # 直接输入到ollama中去，获得模型的回答
        # self.log(title="LLM Response", message=llm_response)

        return [self.extract_sql(llm_response), llm_response, doc_list, ddl_list]

    # 这个是对Vanna里面的函数进行重写，函数的作用是与ollama进行对话
    def submit_prompt(self, prompt, **kwargs) -> str:
        self.log(
            f"Ollama parameters:\n"
            f"model={self.model},\n"
            f"options={self.ollama_options},\n"
            f"keep_alive={self.keep_alive}")
        # self.log(f"Prompt Content:\n{json.dumps(prompt)}")
        # print(prompt)
        response_dict = self.ollama_client.chat(model=self.model,
                                                messages=prompt,
                                                stream=False,
                                                options=self.ollama_options,
                                                keep_alive=self.keep_alive)

        self.log(f"Ollama Response:\n{str(response_dict)}")

        return response_dict['message']['content']

    # 这函数是对Vanna里面的函数进行重写，作用是每次用户提问，生成了sql，查询到数据后，对此次查询进行文字的总结
    def generate_summary(self, question: str, df: pd.DataFrame, **kwargs) -> str:
        """
        **Example:**
        ```python
        vn.generate_summary("What are the top 10 customers by sales?", df)
        ```

        Generate a summary of the results of a SQL query.

        Args:
            question (str): The question that was asked.
            df (pd.DataFrame): The results of the SQL query.

        Returns:
            str: The summary of the results of the SQL query.
        """
        # 这里是在原本的内容上，加上了一些调整
        message_log = [
            self.system_message(
                f"You are a helpful data assistant. The user asked the question: '{question}'\n\n"
                f"The following is a pandas DataFrame with the results of the query: \n"
                f"{df.to_markdown()}\n\n"
            ),
            self.user_message(
                "Briefly summarize the data based on the question that was asked. "
                "Ensure that any ratios or percentages are formatted with the '%' symbol. "
                "Do not respond with any additional explanation beyond the summary. "
                "只需要提供最后的总结" +
                self._response_language()
            ),
        ]

        summary = self.submit_prompt(message_log, **kwargs)

        return summary

    # 这个函数是对Vanna的函数重写，作用是将输入的文本，制作成可以和ollama对话的格式
    def get_sql_prompt(
            self,
            initial_prompt: str,
            question: str,
            question_sql_list: list,
            ddl_list: list,
            doc_list: list,
            **kwargs,
    ):

        if initial_prompt is None:
            initial_prompt = f"You are a {self.dialect} expert. " + \
                             "Please help to generate a SQL query to answer the question. Your response should ONLY be based on the given context and follow the response guidelines and format instructions. "

        initial_prompt = self.add_ddl_to_prompt(
            initial_prompt, ddl_list, max_tokens=32768
        )

        initial_prompt = self.add_documentation_to_prompt(
            initial_prompt, doc_list, max_tokens=32768
        )

        initial_prompt += (

            "===Response Guidelines \n"
            "1. You are a database and knowledge assistant. You can answer data-related questions using the tables in the database, or answer general questions using your knowledge. Before answering, assess whether the question needs a database query. \n"
            "2. For conditions,  treat date fields as timestamps.\n"
            "3. If the provided context is sufficient, please generate a valid SQL query. \n"
            "4. If the provided context is almost sufficient but requires knowledge of a specific string in a particular column, please generate an intermediate SQL query to find the distinct strings in that column. Prepend the query with a comment saying intermediate_sql \n"
            "5. Please use the most relevant table(s). \n"
            "6. If the question has been asked and answered before, please repeat the answer exactly as it was given before. \n"
            f"7. Ensure that the output SQL is {self.dialect}-compliant and executable, and free of syntax errors. \n"
            "生成where的时候，主要不要搞错了字段的内容。\n"
            "SQL查询中将英文列名用AS转换为中文方便理解,中文记得加上引号\n"
            "一步一步的分析问题，生成sql来解决问题给出答案\n"
            "如果这个问题之前已经被问过并且回答过，请准确地重复之前给出的答案。\n"
            "生成 SQL 查询时，无论是做差值、汇总还是对比，请始终避免跨行计算。确保相关字段出现在同一行内进行运算，避免产生不准确的结果。\n"

        )

        message_log = [self.system_message(initial_prompt)]

        message_log.extend(kwargs.get('history', []))  # 提交历史记录

        message_log.append(self.user_message(question))  # 最新的提问

        return message_log

    # 这个函数是对Vanna里面的函数进行重写，作用是将ollama里的大模型的回答进行解析，提取回答里的sql，只获取文本最后的sql，因为那部分的sql才是完整的sql
    def extract_sql(self, llm_response):

        # Remove ollama-generated extra characters
        llm_response = llm_response.replace("\\_", "_")
        llm_response = llm_response.replace("\\", "")
        select_with = re.finditer(r'(?<=```sql\n)((.|\n)*?)(?=;|\[|```)',
                                  llm_response,
                                  re.IGNORECASE | re.DOTALL)

        text = ""
        if select_with:

            for match in select_with:
                text = match.group(0)
            self.log(
                f"Output from LLM: {llm_response} \nExtracted SQL: {text}")
            return text
        else:
            return ""

    # 这个函数是自定义函数，这个用于将markdown格式的文本，制作成文档，包括设置文档标题等内容
    def create_word_from_markdown(self, markdown_text):
        """
        将 Markdown 文本内容解析为 Word 文档。
        :param markdown_text: 大模型生成的 Markdown 文本
        :return: Word 文档对象
        """
        document = Document()
        markdown_text = markdown_text.replace("markdown", "")
        markdown_text = markdown_text.replace("```", "")
        lines = markdown_text.split("\n")

        for line in lines:
            if line.startswith("# "):  # 一级标题
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):  # 二级标题
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):  # 二级标题
                document.add_heading(line[4:], level=3)
            elif line.startswith("#### "):  # 二级标题
                document.add_heading(line[5:], level=4)
            elif line.startswith("##### "):  # 二级标题
                document.add_heading(line[6:], level=5)
            elif line.startswith("- "):  # 列表项
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():  # 普通段落
                document.add_paragraph(line.strip())

        return document

    # 这个函数是自定义函数，这个和get_sql_prompt函数的作用是一样的，专门用于生成word这一功能，制作可以直接和ollama大模型对话的格式
    def generate_message_log(self, history_for_word):
        """
        根据历史对话记录生成 message_log。
        :param history_for_word: 用户选择的历史对话记录（列表）
        :return: 用于生成 Word 文档的 message_log
        """
        # 定义初始的系统消息，这里是对大模型角色的定义
        message_log = [
            self.system_message("你是一个熟练的财务分析师。根据以下对话历史记录，创建一个详细的 Word 报告，格式为 Markdown，报告包含以下部分：\n\n"
                                "1. **前言**：总结查询的背景和目的。\n"
                                "2. **查询内容**：详细列出用户的具体查询项。\n"
                                "3. **数据展示**：展示查询到的数据，包括时间变化趋势。\n"
                                "4. **深入分析**：对数据进行深入分析，例如资产、负债和所有者权益的变化。\n"
                                "5. **结论与建议**：总结发现并提出可操作的建议。")
        ]

        # 将历史记录格式化为用户和助手的对话，这个history_for_word是前端传过来的，用于制作生成文档的prompt
        for i in range(0, len(history_for_word), 2):  # 假设每两个元素是用户提问和大模型回答
            user_message = history_for_word[i]  # 用户提问
            assistant_message = history_for_word[i + 1] if i + 1 < len(history_for_word) else ""

            message_log.append(self.user_message(user_message))
            message_log.append(self.assistant_message(assistant_message))
        # 下面这个就是跟大模型提问的模板
        message_log.append(self.user_message(f"""
### 生成报告要求
你是一个财务分析专家，基于以下对话历史记录，生成一个详细的财务分析报告，格式为 Markdown，报告包含以下部分：

1. **前言**：总结查询的背景和目的。
2. **查询内容**：详细列出用户请求的查询项。
3. **数据展示**：展示查询到的数据，包含时间范围内的变化情况。
4. **深入分析**：对查询的数据进行详细分析.
5. **结论与建议**：总结数据的发现，并给出可操作的建议。

请严格按照这些部分进行总结，不要附加其他额外的解释或信息。输出格式应如下：

```
# 前言
[生成的前言内容]

# 查询内容
[生成的查询内容]

# 数据展示
[展示的数据内容]

# 深入分析
[生成的深入分析内容]

# 结论与建议
[生成的结论与建议内容]
```

请确保：
- 不要有额外的解释或多余的内容。
- 按照要求生成每一部分内容，不要漏掉任何部分。
- 输出必须符合 Markdown 格式，方便生成 Word 文档。
    """))
        return message_log

    # 这个是自定义函数，作用是专门为生成PPT这个功能制作和ollama大模型对话的格式
    def generate_message_log_ppt(self, history_for_word):
        """
        根据历史对话记录生成 message_log。
        :param history_for_word: 用户选择的历史对话记录（列表）
        :return: 用于生成 Word 文档的 message_log
        """
        # 定义初始的系统消息，这里是对大模型角色的定义
        message_log = [
            self.system_message("你是一个熟练的财务PPT生成专家。根据以下对话历史记录，创建一个详细的 PPT，格式为 json，PPT包含以下部分：\n\n"
                                "1. **前言**：总结查询的背景和目的。\n"
                                "2. **查询内容**：详细列出用户的具体查询项。\n"
                                "3. **数据展示**：展示查询到的数据，包括时间变化趋势。\n"
                                "4. **深入分析**：对数据进行深入分析，例如资产、负债和所有者权益的变化。\n"
                                "5. **结论与建议**：总结发现并提出可操作的建议。")
        ]

        # 将历史记录格式化为用户和助手的对话
        for i in range(0, len(history_for_word), 2):  # 假设每两个元素是用户提问和大模型回答
            user_message = history_for_word[i]  # 用户提问
            assistant_message = history_for_word[i + 1] if i + 1 < len(history_for_word) else ""

            message_log.append(self.user_message(user_message))
            message_log.append(self.assistant_message(assistant_message))
        message_log.append(self.user_message(
            """
### 生成PPT要求  
你是一位资深的财务 PPT 生成专家，基于以下对话历史记录，生成一个详细的财务分析 PPT，格式为 **JSON**。PPT 的内容需包括以下部分：  

1. **前言**  
   - 总结查询的背景和目的。  
2. **查询内容**  
   - 列出用户请求的具体查询项。  
3. **数据展示**  
   - 展示历史记录中查询到的数据，包含时间范围内的变化情况（如趋势、对比等）。  
   - 如果历史记录中有表格数据，请将其格式化为易于展示的方式（如简要摘要或数据点）。  
4. **深入分析**  
   - 基于展示的数据进行详细分析，包括发现的模式、潜在问题或其他相关洞察。  
5. **结论与建议**  
   - 总结主要发现，并提供具体、可操作的建议。  

### 输出要求  
生成的 JSON 必须符合以下格式：  

```json  
{  
    "title": "演示文稿标题",  
    "slides": [  
        {  
            "heading": "第一张幻灯片的标题",  
            "bullet_points": [  
                "第一项要点",  
                [  
                    "子要点 1",  
                    "子要点 2"  
                ],  
                "第二项要点"  
            ],  
            "key_message": ""  
        },  
        {  
            "heading": "第二张幻灯片的标题",  
            "bullet_points": [  
                "第一项要点",  
                "第二项要点",  
                "第三项要点"  
            ],  
            "key_message": "这张幻灯片传达的关键信息"  
        },  
        {  
            "heading": "描述逐步过程的幻灯片",  
            "bullet_points": [  
                ">> 过程的第一步（以特殊标记 >> 开头）",  
                ">> 第二步（以 >> 开头）",  
                ">> 第三步"  
            ],  
            "key_message": ""  
        },  
        {  
            "heading": "双栏布局的幻灯片（用于并排对比或对照两个相关的概念，例如优缺点、优势与风险、旧方法与现代方法等）",  
            "bullet_points": [  
                {  
                    "heading": "左栏标题",  
                    "bullet_points": [  
                        "第一项要点",  
                        "第二项要点",  
                        "第三项要点"  
                    ]  
                },  
                {  
                    "heading": "右栏标题",  
                    "bullet_points": [  
                        "第一项要点",  
                        "第二项要点",  
                        "第三项要点"  
                    ]  
                }  
            ],  
            "key_message": ""  
        }  
    ]  
}  
"""
        ))
        return message_log

    # 这个是对Vanna里面函数的重写，用于连接数据库，重点：已经适配了人大金仓数据库
    def connect_to_postgres(
            self,
            host: str = None,
            dbname: str = None,
            user: str = None,
            password: str = None,
            port: int = None,
            **kwargs
    ):

        try:
            import psycopg2
            import psycopg2.extras
        except ImportError:
            raise RuntimeError(
                "You need to install required dependencies to execute this method,"
                " run command: \npip install vanna[postgres]"
            )

        if not host:
            host = os.getenv("HOST")

        if not host:
            raise RuntimeError("Please set your postgres host")

        if not dbname:
            dbname = os.getenv("DATABASE")

        if not dbname:
            raise RuntimeError("Please set your postgres database")

        if not user:
            user = os.getenv("PG_USER")

        if not user:
            raise RuntimeError("Please set your postgres user")

        if not password:
            password = os.getenv("PASSWORD")

        if not password:
            raise RuntimeError("Please set your postgres password")

        if not port:
            port = os.getenv("PORT")

        if not port:
            raise RuntimeError("Please set your postgres port")

        conn = None

        try:
            conn = psycopg2.connect(
                host=host,
                dbname=dbname,
                user=user,
                password=password,
                port=port,
                **kwargs
            )
        except psycopg2.Error as e:
            raise RuntimeError(e)

        def connect_to_db():
            return psycopg2.connect(host=host, dbname=dbname,
                                    user=user, password=password, port=port, **kwargs)

        def run_sql_postgres(sql: str) -> Union[pd.DataFrame, None]:
            conn = None
            try:
                conn = connect_to_db()  # Initial connection attempt
                conn.set_client_encoding('UTF8')
                cs = conn.cursor()
                cs.execute(sql)
                results = cs.fetchall()

                # Create a pandas dataframe from the results
                df = pd.DataFrame(results, columns=[desc[0] for desc in cs.description])
                return df

            except psycopg2.InterfaceError as e:
                # Attempt to reconnect and retry the operation
                if conn:
                    conn.close()  # Ensure any existing connection is closed
                conn = connect_to_db()
                cs = conn.cursor()
                cs.execute(sql)
                results = cs.fetchall()

                # Create a pandas dataframe from the results
                df = pd.DataFrame(results, columns=[desc[0] for desc in cs.description])
                return df

            except psycopg2.Error as e:
                if conn:
                    conn.rollback()
                    raise RuntimeError(e)

            except Exception as e:
                conn.rollback()
                raise e

        self.dialect = "PostgreSQL"
        self.run_sql_is_set = True
        self.run_sql = run_sql_postgres

    # 这个是对Vanna里面函数的重写，用于让大模型生成plot代码，用于生成图表
    def generate_plotly_code(
            self, question: str = None, sql: str = None, df_metadata: str = None, **kwargs
    ) -> str:
        if question is not None:
            system_msg = f"The following is a pandas DataFrame that contains the results of the query that answers the question the user asked: '{question}'"
        else:
            system_msg = "The following is a pandas DataFrame "

        if sql is not None:
            system_msg += f"\n\nThe DataFrame was produced using this query: {sql}\n\n"

        system_msg += f"The following is information about the resulting pandas DataFrame 'df': \n{df_metadata}"

        message_log = [
            self.system_message(system_msg),
            self.user_message(
                "Can you generate the Python plotly code to chart the results of the dataframe? Assume the data is in a pandas dataframe called 'df'. If there is only one value in the dataframe, use an Indicator. Respond with only Python code and Chinese. Do not answer with any explanations -- just the code。"
            ),
        ]

        plotly_code = self.submit_prompt(message_log, kwargs=kwargs)

        return self._sanitize_plotly_code(self._extract_python_code(plotly_code))

    # 这个函数是对Vanna里面函数的重写，用于对大模型的回答里，提取python代码，这代码是生成图表的代码。同样只需要提取最终的代码，所以python_code[-1]是取的-1
    def _extract_python_code(self, markdown_string: str) -> str:
        # Regex pattern to match Python code blocks
        pattern = r"```[\w\s]*python\n([\s\S]*?)```|```([\s\S]*?)```"

        # Find all matches in the markdown string
        matches = re.findall(pattern, markdown_string, re.IGNORECASE)

        # Extract the Python code from the matches
        python_code = []
        for match in matches:
            python = match[0] if match[0] else match[1]
            python_code.append(python.strip())

        if len(python_code) == 0:
            return markdown_string

        return python_code[-1]


# 这个是创建嵌入模型函数的对象
embedding_function = OllamaEmbeddingFunction(model_name="bge-m3:latest")
# "embedding_function":embedding_function,
vn = MyVanna(config={"embedding_function": embedding_function, 'model': 'qwen:0.5b', "n_results_sql": 4,
                     "n_results_documentation": 1, "n_results_ddl": 1, "ollama_host": "http://192.168.30.130:11434",
                     'options': {"num_ctx": 2000}, "max_tokens": 32768, "language": "chinese", "initial_prompt":
                         f"你是一个数据库和知识系统的助手，你可以根据数据库中的表来回答数据相关问题，也可以通过你的知识回答一般性问题。\n"
                         f"请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL；\n"
                         f"如果不需要，请直接回答\n"
                         f"日期字段的属性是timestamp,\n"
                         "生成where的时候，主要不要搞错了字段的内容\n"
                         "SQL查询中将英文列名用AS转换为中文方便理解,中文记得加上引号\n"
                         "一步一步的分析问题，生成sql来解决问题给出答案\n"
                         "生成 SQL 查询时，无论是做差值、汇总还是对比，请始终避免跨行计算。确保相关字段出现在同一行内进行运算，避免产生不准确的结果。\n"})
# 连接数据库，数据库部署在阿里云上
vn.connect_to_postgres(host='192.168.30.130', dbname='postgres', user='test', password='test', port=5432)
# 这些是用于程序运行中的变量
history = []
doc_list_tmp = []
ddl_list_tmp = []
data_rag = {}
# region 对数据表进行加载
# 下面是对每张数据表进行背景知识的构建，对每张表进行详细的解释，然后存如data_rag字典里去，之后会根据chromdb的检索结果，来从这里面取特定的一张表，输入到大模型对话里面去
# 后面是对8张表进行一样的处理
# 表cwgk_report_vb_cwys031

# #twc_km_bal_m
data_rag['twb_ygl0078'] = []
data_rag['twb_ygl0078'].append(
    """
    -- public.twb_ygl0078 definition
    
    -- Drop table
    
    -- DROP TABLE public.twb_ygl0078;
    
    CREATE TABLE public.twb_ygl0078 (
        organ_name varchar(255) NULL, -- 公司名称
        tdate timestamp NULL, -- 日期
        mat_ind_name varchar(255) NULL, -- 指标名
        unit_name varchar(255) NULL, -- 单位
        curr_y_budget float8 NULL, -- 月预算
        curr_m_pe_num float8 NULL -- 月金额
    );
    
    -- Column comments
    
    COMMENT ON COLUMN public.twb_ygl0078.organ_name IS '公司名称';
    COMMENT ON COLUMN public.twb_ygl0078.tdate IS '日期';
    COMMENT ON COLUMN public.twb_ygl0078.mat_ind_name IS '指标名';
    COMMENT ON COLUMN public.twb_ygl0078.unit_name IS '单位';
    COMMENT ON COLUMN public.twb_ygl0078.curr_y_budget IS '月预算';
    COMMENT ON COLUMN public.twb_ygl0078.curr_m_pe_num IS '月金额';
    """
)
data_rag['twb_ygl0078'].append(
    """
    1.organ_name: 公司名称，用于表示公司。
    2.tdate: 录入日期，记录数据的录入时间。
    3.mat_ind_name: 指标名称，表示财务的指标种类。
    4.unit_name: 单位，金额单位。
    5.curr_y_budget: 月预算，指每个月预期的收入金额。
    6.curr_m_pe_num: 科目名称，描述财务科目。\n"""
)
mat_ind_name_5 = vn.run_sql("SELECT DISTINCT mat_ind_name FROM twb_ygl0078 WHERE mat_ind_name IS NOT NULL ")
organ_name_5 = vn.run_sql("SELECT DISTINCT organ_name FROM twb_ygl0078 WHERE organ_name IS NOT NULL ")

mat_ind_name_5_list = [item for item in mat_ind_name_5['mat_ind_name'].tolist() if item is not None]
print(mat_ind_name_5_list)
organ_name_5_list = [item for item in organ_name_5['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
mat_ind_name_5_str = ', '.join(mat_ind_name_5_list)
organ_name_5_str = ', '.join((organ_name_5_list))
data_rag['twb_ygl0078'].append(
    """
    twb_ygl0078表
    """ +
    f"""**关键字段**:
1. **mat_ind_name** (科目名):
   - **固定取值列表**: 
     - [mat_ind_name_5_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_5_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **item_name** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=mat_ind_name_5_list, name='twb_ygl0078')

tmp_question = "，请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL，如果不需要，请直接回答，如果需要生成sql，生成一次sql，请严格按照system内容的固定取值列表来生成sql，注意一定要在sql中给出结果。"
# endregion

import os.path
import sqlite3
import time
from functools import lru_cache
from typing import List

from gradio_client import Client
from difflib import SequenceMatcher

db_path = "/code/kotaemon/ktem_app_data/user_data/sql.db"


# region 这个是用于实现知识库问答的函数，目前是直接通过api的方式，对kotaemon的功能进行调用，实际上应该把kotaemon的rag代码提取出来，放到这里
@lru_cache(1)  # 这个是自定义函数
def get_all_file_ids(db_path: str, indices: List[int] | None = None):
    assert os.path.isfile(db_path)

    file_ids: List[str] = []

    conn = sqlite3.connect(db_path)

    try:
        c = conn.cursor()
        if indices is None:
            # Get all fileIndex
            c.execute("SELECT * FROM ktem__index")
            indices = [each[0] for each in c.fetchall()]

        # For each index get all file_ids
        for i in indices:
            table_name = f"index__{i}__source"
            c.execute(
                f"SELECT * FROM {table_name}",
            )
            file_ids += [each[0] for each in c.fetchall()]
    finally:
        conn.close()

    return file_ids


def find_differences(str1, str2):
    matcher = SequenceMatcher(None, str1, str2)
    differences = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != 'equal':
            differences.append(str2[j1:j2])
    return "".join(differences)


def chat_rag(question):
    file_ids = get_all_file_ids(db_path)

    print(f"File_ids: {file_ids}")

    client = Client("http://localhost:7860/")

    job = client.submit(
        [
            [
                question + ",用中问回答",
                None,
            ]
        ],
        "select",
        file_ids,
        api_name="/chat_fn",
    )

    resps = ['']
    full_response = ""

    # 逐步收集响应并保存到文件
    for o in job:
        # 检查是否有新的内容返回
        if 'Thinking' not in o[0][0][1]:
            if len(resps[-1]) != len(o[0][0][1]):
                content = find_differences(resps[-1], o[0][0][1])
                resps.append(o[0][0][1])
                full_response += content

            else:
                break

    return full_response


# endregion
# region 这个是用于生成ppt的函数，是借鉴于slide_deck项目里面的ppt生成方法
def generate_slide_deck(json_str: str):
    """
    Create a slide deck.

    :param json_str: The content in *valid* JSON format.
    """

    path = "financial_report.pptx"
    pptx_template = r"./pptx_templates/Blank.pptx"

    try:
        ppt = pptx_helper.generate_powerpoint_presentation(
            json_str,
            slides_template=pptx_template,
            output_file_path=path
        )  # 生成幻灯片文件
    except ValueError:

        ppt = pptx_helper.generate_powerpoint_presentation(
            text_helper.fix_malformed_json(json_str),
            slides_template=pptx_template,
            output_file_path=path
        )  # 生成幻灯片文件
    return ppt


# endregion
def requires_cache(fields):  # 这个函数和缓存进行搭配
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            id = request.args.get('id')

            if id is None:
                return jsonify({"type": "error", "error": "No id provided"})

            for field in fields:
                if cache.get(id=id, field=field) is None:
                    return jsonify({"type": "error", "error": f"No {field} found"})

            field_values = {field: cache.get(id=id, field=field) for field in fields}

            # Add the id to the field_values
            field_values['id'] = id

            return f(*args, **field_values, **kwargs)

        return decorated

    return decorator


# 下面是一些函数的定义，主要是和前端进行交互，下面一一进行介绍
# region 这个函数是用于界面刚刷新的时候，生成一些问题，目前点击问题是没有任何效果的
@app.route('/api/v0/generate_questions', methods=['GET'])
def generate_questions():
    # vn.generate_questions()
    return jsonify({
        "type": "question_list",
        "questions": ["告诉我H公司在23年4月份，营业收入的本月完成总金额是多少", "告诉我D公司合并在22年11月份，本月的综合收益总额是多少"],
        "header": "这里有一些您可能想查询的数据:"
    })


# endregion
# region 这个函数是和前端交互，生成sql答案的，这里面也包含了sql自纠正
@app.route('/api/v0/generate_sql', methods=['GET'])
def generate_sql():
    question = flask.request.args.get('question')
    former_doc_list = flask.request.args.get("former_doc_list")
    mode = flask.request.args.get("mode_web")
    id = cache.generate_id(question=question)

    global doc_list_tmp
    global ddl_list_tmp
    # question=question+"。你是一个数据库和知识系统的助手，你可以根据数据库中的表来回答数据相关问题，也可以通过你的知识回答一般性问题。请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL,确认好表的字段，不要搞错了；如果不需要，请直接回答,在生成sql时将WHERE子句中使用的所有字段添加到 SELECT中，生成的sql，用like代替等于号，同时like要搭配%符号，日期字段的属性是timestamp"
    if question is None:
        return jsonify({"type": "error", "error": "No question provided"})

    if mode == "file":
        return jsonify(
            {
                "type": "sql",
                "id": id,
                "sql_que": "",
                "text": chat_rag(question),  # 文档查询的
            })

    else:
        if former_doc_list == "true":

            response = vn.generate_sql(question=question, allow_llm_to_see_data=True, history=history,
                                       doc_list=doc_list_tmp, ddl_list=ddl_list_tmp)
        else:
            # question+="，根据上下文判断是否需要生成sql，如果需要生成sql，就只生成一次sql，在sql中得出答案"
            response = vn.generate_sql(question=question, allow_llm_to_see_data=True, history=history, doc_list=[])
            doc_list_tmp = response[2]
            ddl_list_tmp = response[3]
        if response[0] != "":
            cache.set(id=id, field='question', value=question)
            cache.set(id=id, field='sql', value=response[0])
        history.extend([{"role": "user", "content": question}, {"role": "assistant", "content": response[0]}])

        return jsonify(
            {
                "type": "sql",
                "id": id,
                "sql_que": response[0],
                "text": response[1],
            })


# endregion
# region 这个函数是用于和前端交互，生成word文档的
@app.route('/api/v0/generate_word', methods=['POST'])
def generate_word():
    data = flask.request.get_json()

    # 获取 reportList，确保后端正确处理传来的数据
    history_for_word = data.get('reportList')
    message_log = vn.generate_message_log(history_for_word)
    print(message_log)
    word_text = vn.submit_prompt(message_log)
    word_text = vn.create_word_from_markdown(word_text)

    output = BytesIO()
    # output.write(word_text.encode('utf-8'))  # 转为字节流
    word_text.save(output)
    output.seek(0)  # 重置流的位置，以确保从头读取

    # 返回 Word 文档作为文件
    return Response(
        output.getvalue(),  # 返回字节数据
        mimetype="application/msword",  # 或 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        headers={"Content-Disposition": "attachment; filename=financial_report.docx"}
    )


# endregion
# region 这个函数和前端交互，用于生成ppt
@app.route('/api/v0/generate_PPT', methods=['POST'])
def generate_PPT():
    data = flask.request.get_json()

    # 获取 reportList，确保后端正确处理传来的数据
    history_for_word = data.get('reportList')
    message_log = vn.generate_message_log_ppt(history_for_word)
    print(message_log)
    response = vn.submit_prompt(message_log)
    response_cleaned = text_helper.get_clean_json(response)
    generate_ppt = generate_slide_deck(response_cleaned)
    # presentation = pptx.Presentation
    # 将生成的 pptx 文件保存在内存中
    ppt_io = BytesIO()
    generate_ppt.save(ppt_io)
    ppt_io.seek(0)  # 将指针移动到文件开头

    return Response(
        ppt_io.getvalue(),  # 返回字节数据
        mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",  # MIME 类型
        headers={"Content-Disposition": "attachment; filename=generated_presentation.pptx"}  # 设置下载文件名
    )


# endregion

# region 这函数是用于执行sql语句
@app.route('/api/v0/run_sql', methods=['GET'])
@requires_cache(['sql'])
def run_sql(id: os, sql: os):
    # def run_sql():
    try:
        # id = flask.request.args.get('id')
        # sql = flask.request.args.get('sql')
        df = vn.run_sql(sql=sql)
        for col in df.columns:
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                df[col] = df[col].dt.strftime('%Y-%m-%d')  # 修改日期格式为 "YYYY-MM-DD"
            elif pd.api.types.is_object_dtype(df[col]):
                try:
                    temp_col = pd.to_datetime(df[col], errors='coerce')  # 将字符串列转为日期，无法转的变为 NaT
                    # 如果转换后有有效的日期，则格式化，否则保持原样
                    if temp_col.notna().any():
                        df[col] = temp_col.dt.strftime('%Y-%m-%d')  # 格式化有效日期
                except Exception as e:
                    print(f"Error in converting column {col}: {e}")

        cache.set(id=id, field='df', value=df)
    except Exception as e:
        error_message = "SQL执行错误,重新生成:" + str(e)
        print(error_message)
        return jsonify({"type": "error", "error": error_message})
    # history[-1]["content"]+="\n\n查询结果如下：\n\n"+df.to_markdown(index=False)
    # history.extend([{"role": "user", "content":"\n\n查询结果如下：\n\n"+df.to_markdown(index=False)},
    #                 {"role": "assistant", "content": "好的"}])
    return jsonify(
        {
            "type": "df",
            "id": id,
            "df": df.to_json(orient='records'),
        })

    # except Exception as e:
    #     return jsonify({"type": "error", "error": os(e)})


# endregion
# region 这个函数是用于前端进行csv的下载
@app.route('/api/v0/download_csv', methods=['GET'])
@requires_cache(['df'])
def download_csv(id: os, df):
    output = BytesIO()
    # 将 CSV 编码为 GBK 并写入字节流
    df.to_csv(output, encoding="GBK", index=False)
    output.seek(0)  # 重置流的位置，以确保从头读取

    return Response(
        output.getvalue(),  # 返回字节数据
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={id}.csv"}
    )


# endregion
# region 这个函数是用于前端图表的生成
@app.route('/api/v0/generate_plotly_figure', methods=['GET'])
@requires_cache(['df', 'sql'])
def generate_plotly_figure(id: os, df, sql):
    question = flask.request.args.get('question')
    try:

        code = vn.generate_plotly_code(question=question, sql=sql,
                                       df_metadata=f"Running df.dtypes gives:\n {df.dtypes}")
        print(code)
        summery = vn.generate_summary(question, df)
        fig = vn.get_plotly_figure(plotly_code=code, df=df, dark_mode=False)
        fig_json = fig.to_json()

        cache.set(id=id, field='fig_json', value=fig_json)

        return jsonify(
            {
                "type": "plotly_figure",
                "id": id,
                "fig": fig_json,
                "summary": summery
            })
    except Exception as e:
        # Print the stack trace
        import traceback
        traceback.print_exc()

        return jsonify({"type": "error", "error": os(e)})


# endregion
# region 这个代码是用于获取训练数据的函数，这个是Vanna原本的函数，不过现在已经不需要这个函数了
@app.route('/api/v0/get_training_data', methods=['GET'])
def get_training_data():
    df = vn.get_training_data()

    return jsonify(
        {
            "type": "df",
            "id": "training_data",
            "df": df.head(25).to_json(orient='records'),
        })


# endregion
# region 这个函数是移除训练的数据，也已经不需要了
@app.route('/api/v0/remove_training_data', methods=['POST'])
def remove_training_data():
    # Get id from the JSON body
    id = flask.request.json.get('id')

    if id is None:
        return jsonify({"type": "error", "error": "No id provided"})

    if vn.remove_training_data(id=id):
        return jsonify({"success": True})
    else:
        return jsonify({"type": "error", "error": "Couldn't remove training data"})


# endregion
# region 这个函数是用于将数据进行训练，但这个已经不需要使用了
@app.route('/api/v0/train', methods=['POST'])
def add_training_data():
    question = flask.request.json.get('question')
    sql = flask.request.json.get('sql')
    ddl = flask.request.json.get('ddl')
    documentation = flask.request.json.get('documentation')

    try:
        id = vn.train(question=question, sql=sql, ddl=ddl, documentation=documentation)

        return jsonify({"id": id})
    except Exception as e:
        print("TRAINING ERROR", e)
        return jsonify({"type": "error", "error": os(e)})


# endregion
# region 同样这个也不需要使用了
@app.route('/api/v0/generate_followup_questions', methods=['GET'])
@requires_cache(['df', 'question', 'sql'])
def generate_followup_questions(id: os, df, question, sql):
    followup_questions = vn.generate_followup_questions(question=question, sql=sql, df=df)

    cache.set(id=id, field='followup_questions', value=followup_questions)

    return jsonify(
        {
            "type": "question_list",
            "id": id,
            "questions": followup_questions,
            "header": "Here are some followup questions you can ask:"
        })


# endregion
# region同样也不需要这个函数
@app.route('/api/v0/load_question', methods=['GET'])
@requires_cache(['question', 'sql', 'df', 'fig_json', 'followup_questions'])
def load_question(id: os, question, sql, df, fig_json, followup_questions):
    try:
        return jsonify(
            {
                "type": "question_cache",
                "id": id,
                "question": question,
                "sql": sql,
                "df": df.head(10).to_json(orient='records'),
                "fig": fig_json,
                "followup_questions": followup_questions,
            })

    except Exception as e:
        return jsonify({"type": "error", "error": os(e)})


# endregion
# region这个是用于在界面显示历史会话，但现在暂时没使用这个函数
@app.route('/api/v0/get_question_history', methods=['GET'])
def get_question_history():
    return jsonify({"type": "question_history", "questions": cache.get_all(field_list=['question'])})


# endregion
@app.route('/')
def root():
    return app.send_static_file('index.html')


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=False)
    # print(vn.submit_prompt(model="qwq:32b",stream=False,prompt=[{"role": "user","content": "你知道sql如何写吗，假如你是一个财务总监"}]))
    # npm create vite@latest my-svelte-app
    # npm install svelte@5.0.0
