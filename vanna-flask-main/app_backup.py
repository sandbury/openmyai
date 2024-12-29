import pandas as pd  # pandas 是数据分析中不可或缺的工具,用于对sql查询的表格进行处理
import re  # re 是 Python 内置的正则表达式模块，用于对大模型的输出进行文本匹配
from functools import wraps  # 装饰器
from flask import Flask, jsonify, Response, request, redirect, url_for  # 导入flask的快捷库
import flask
import os
from cache import MemoryCache  # 利用字典来作为临时的内存存储
from vanna.chromadb import ChromaDB_VectorStore  # 导入Vanna里设置好的chromdb文件
from vanna.ollama import Ollama  # 导入Vanna配置好的ollama交互文件
from flask import Flask, request, send_file
from io import BytesIO  # BytesIO 是一个内存中的字节流对象，在这里用于把文件存在内存中，然后供前端下载
from docx import Document  # 用于生成文档的库
import json
from typing import List, Tuple, Union  # 导入一些数据类型
import ollama
from helpers import pptx_helper, text_helper  # 从helpers文件夹中导入生成ppt的函数

app = Flask(__name__, static_url_path='')

# SETUP
cache = MemoryCache()

# from vanna.local import LocalContext_OpenAI
# vn = LocalContext_OpenAI()

# from vanna.remote import VannaDefault
# vn = VannaDefault(model=os.environ['VANNA_MODEL'], api_key=os.environ['VANNA_API_KEY'])

from typing import List
from chromadb.utils.embedding_functions import EmbeddingFunction


# 自定义函数：用于在Vanna类里面使用自己定义的嵌入模型
class OllamaEmbeddingFunction(EmbeddingFunction[List[str]]):
    def __init__(self, model_name="nomic-embed-text:latest"):
        self.model_name = model_name

    def __call__(self, input: List[str]) -> List[List[float]]:
        embeddings = []
        for text in input:
            # 使用 ollama.embeddings 获取嵌入向量
            response = ollama.embeddings(
                model=self.model_name,
                prompt=text
            )
            embedding = response['embedding']  # 获取嵌入向量
            embeddings.append(embedding)

        return embeddings


collections = {}  # 这个是用于装载定义的数据库背景知识，包括数据表的ddl，对字段的解释等内容。


# 这个是自定义一个叫Vanna的类，集成ChromaDB_VectorStore和Ollama这两个类，用于本地部署数据库和大模型
class MyVanna(ChromaDB_VectorStore, Ollama):
    def __init__(self, config=None):
        ChromaDB_VectorStore.__init__(self, config=config)
        Ollama.__init__(self, config=config)

    # 这个函数是自定义的，只需要执行一次就行，下次再启动项目的时候，把if len(text)>1:这后面的代码注释掉
    # 这个函数是将collections里面装载的内容存到本地chromadb数据库中，因为第一次要把内容向量化，然后再存进数据库，所以第一次运行比较久
    # 第二次运行就没必要在存数据库了，所以就把if len(text)>1:这后面的代码注释掉（这种方式不太方便，应该优化）
    def train_yu(
            self,
            text=[],
            name=""
    ):
        if name:
            collections[name] = self.chroma_client.get_or_create_collection(
                name=name,
            )
        # if len(text)>1:
        #     for i, d in enumerate(text):
        #         response = ollama.embeddings(model="bge-m3:latest", prompt=d)
        #         embedding = response["embedding"]
        #         collections[name].add(
        #             ids=[str(i)],
        #             embeddings=[embedding],
        #             documents=[d]
        #         )

    # 这个函数是自定义函数，用于将用户提问的问题，和数据库中每张表的固定的字段内容进行匹配，然后筛选出哪张表是最适合用户提问的问题的。
    def get_similar_text_yu(self, question):
        best_score = 1000
        best_data = []
        best_name = ""
        ddl_list = []
        doucment_list = []
        response = ollama.embeddings(
            prompt=question,
            model="bge-m3:latest"
        )

        for i in collections.keys():
            if i in question:
                results = collections[i].query(
                    query_embeddings=[response["embedding"]],
                    n_results=20,
                )
                data = results['documents'][0]
                score = results['distances'][0]
                best_name = i
                best_data = data
                break
            else:
                results = collections[i].query(
                    query_embeddings=[response["embedding"]],
                    n_results=20,
                )
                data = results['documents'][0]
                score = results['distances'][0]
                if score[0] < best_score:
                    best_name = i
                    best_score = score[0]
                    best_data = data
        # ddl_list.append(data_rag[best_name][0])
        # doucment_list.append(data_rag[best_name][1])
        if best_name == 'cwgk_report_vb_cwys031':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("item_names_str", ",".join(best_data))
        elif best_name == 'sap_yzb0003':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("proj_name_2_str", ",".join(best_data))
        elif best_name == 'twb_yzb0102':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("asset_desc_4_str", ",".join(best_data))
        elif best_name == 'twc_km_bal_m':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("item_name_5_str", ",".join(best_data))
        elif best_name == 'twc_profit_state_m':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("proj_name_6_str", ",".join(best_data))
        elif best_name == 'twc_thermal_power_operate_m':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("proj_name_7_str", ",".join(best_data))
        elif best_name == 'twc_ygl0060':
            tmp = data_rag[best_name][2]
            tmp = tmp.replace("proj_name_8_str", ",".join(best_data))
        elif best_name == 'twb_ygl0078':
            tmp = data_rag[best_name][2]
        # tmp=data_rag[best_name][2]

        doucment_list.append(data_rag[best_name][0] + "\n" + data_rag[best_name][1] + "\n" + tmp)

        return [], doucment_list, tmp

    # 这个函数是对Vanna里面的函数重写，修改了一些流程
    def generate_sql(self, question: str, allow_llm_to_see_data=False, doc_list=[], ddl_list=[], **kwargs) -> list:

        if self.config is not None:
            initial_prompt = self.config.get("initial_prompt", None)
        else:
            initial_prompt = None

        # question_sql_list = self.get_similar_question_sql(question, **kwargs)
        question_sql_list = []  # 这个变量暂时不用了，暂时不需要提供  问题-sql答案  这样的文本了

        # ddl_list = self.get_related_ddl(question, **kwargs)
        context_tmp = ""
        # ddl_list=[]
        if len(doc_list) == 0:  # doc_list，ddl_list，用来实现sql自纠错机制
            ddl_list, doc_list, context_tmp = self.get_similar_text_yu(question)  # doc_list是system的内容
        # 在这里加上提示词,规范输出
        template_input = f"""
        
---

### 规则说明
1. **判断是否需要查询数据库**：
   - 如果问题涉及数据查询或计算，请先一步步分析问题，再生成 SQL。
   - 如果问题无需查询数据库，请直接回答用户问题。
2. **SQL 生成规则**：
   - 表名、字段和条件必须严格按照 system 提供的固定取值范围。
   - 使用 `WITH` 结构组织查询，以提高可读性和扩展性。
   - 确保查询在 SQL 中给出具体的结果。
   - **无论查询还是计算，若表中有单位字段，必须在最终结果中包含单位字段**。
3. **单位字段的特殊规则**：
   - 如果其他字段使用了聚合函数（如 `SUM`、`AVG` 等），单位字段应加入 `GROUP BY` 子句。
   - 如果计算结果与单位字段无直接关联（如跨列计算），结果表中仍需保留单位字段，作为辅助信息。
4. **特殊注意**：
   - 避免跨行计算，所有计算仅在单行内完成。
   - 日期字段类型为 `timestamp`，查询中应考虑日期格式兼容。
   - **比率计算时不再乘以 100，直接输出比率值**。
   - 如果查询中的字段发生别名变化（例如 AS 用法），应确保后续 SELECT 使用正确的别名字段，而不是原字段名
---

### 用户提问
{question}

---

### 输出要求
根据用户提问和规则说明：
1. 判断问题是否需要查询数据库。
   - 如果需要，请一步步分析问题后生成 SQL。
   - 如果不需要，请直接回答问题。
2. **生成 SQL 的结构**：
   - 使用 `WITH` 子句进行逻辑拆分。
   - 确保查询结果输出具体的计算或筛选值。
   - **无论查询还是计算，结果中必须包含单位字段，且不得引发 SQL 语法错误**。
   - **确保 SQL 查询中的公司和日期字段始终按需要进行分组，并避免跨行计算或计算错误**。
   - 如果有聚合函数，单位字段必须加入 `GROUP BY` 子句。
   - 日期字段类型为 `timestamp`，查询中应考虑日期格式兼容。
3. **SQL 示例格式**：
   ```sql
   -- 请将生成的 SQL 放置在此格式中，确保语法正确
   ```
"""
        # question+="，请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL，如果不需要，请直接回答，如果需要生成sql，一步一步的分析问题，生成sql来解决问题给出答案，最好使用with的结构，请严格按照system内容的固定取值列表来生成sql，注意一定要在sql中给出结果。"
        prompt = self.get_sql_prompt(  # 将自己定义的问题，制作成可以直接输入到ollama里去，与大模型交流的格式
            initial_prompt=initial_prompt,
            question=template_input,
            question_sql_list=question_sql_list,
            ddl_list=ddl_list,
            doc_list=doc_list,
            **kwargs,
        )
        self.log(title="SQL Prompt", message=prompt)
        llm_response = self.submit_prompt(prompt, **kwargs)  # 直接输入到ollama中去，获得模型的回答
        # self.log(title="LLM Response", message=llm_response)

        # 下面注释的是，如果大模型有中间sql的生成，要先执行中间sql，但是现在大模型已经不会生成中间sql了，大模型会在回答的末尾，整合sql答案，所以下面的代码不需要了
        # if 'intermediate_sql' in llm_response:
        #     if not allow_llm_to_see_data:
        #         return "The LLM is not allowed to see the data in your database. Your question requires database introspection to generate the necessary SQL. Please set allow_llm_to_see_data=True to enable this."
        #
        #     if allow_llm_to_see_data:
        #         intermediate_sql = self.extract_sql(llm_response)
        #
        #         try:
        #             self.log(title="Running Intermediate SQL", message=intermediate_sql)
        #             df = self.run_sql(intermediate_sql)
        #
        #             prompt = self.get_sql_prompt(
        #                 initial_prompt=initial_prompt,
        #                 question=question,
        #                 question_sql_list=question_sql_list,
        #                 ddl_list=ddl_list,
        #                 doc_list=doc_list+[f"The following is a pandas DataFrame with the results of the intermediate SQL query {intermediate_sql}: \n" + df.to_markdown()],
        #                 **kwargs,
        #             )
        #             # self.log(title="Final SQL Prompt", message=prompt)
        #             llm_response = self.submit_prompt(prompt, **kwargs)
        #             # self.log(title="LLM Response", message=llm_response)
        #         except Exception as e:
        #             return f"Error running intermediate SQL: {e}"

        return [self.extract_sql(llm_response), llm_response, doc_list, ddl_list]

    # 这个是对Vanna里面的函数进行重写，函数的作用是与ollama进行对话
    def submit_prompt(self, prompt, **kwargs) -> str:
        self.log(
            f"Ollama parameters:\n"
            f"model={self.model},\n"
            f"options={self.ollama_options},\n"
            f"keep_alive={self.keep_alive}")
        # self.log(f"Prompt Content:\n{json.dumps(prompt)}")
        # print(prompt)
        response_dict = self.ollama_client.chat(model=self.model,
                                                messages=prompt,
                                                stream=False,
                                                options=self.ollama_options,
                                                keep_alive=self.keep_alive)
        # res=ollama.chat(model="qwq:32b",
        #                 stream=False,
        #                 messages=prompt,
        #                 options={'num_ctx': 6000}
        #                 )

        self.log(f"Ollama Response:\n{str(response_dict)}")

        return response_dict['message']['content']

    # 这函数是对Vanna里面的函数进行重写，作用是每次用户提问，生成了sql，查询到数据后，对此次查询进行文字的总结
    def generate_summary(self, question: str, df: pd.DataFrame, **kwargs) -> str:
        """
        **Example:**
        ```python
        vn.generate_summary("What are the top 10 customers by sales?", df)
        ```

        Generate a summary of the results of a SQL query.

        Args:
            question (str): The question that was asked.
            df (pd.DataFrame): The results of the SQL query.

        Returns:
            str: The summary of the results of the SQL query.
        """
        # 这里是在原本的内容上，加上了一些调整
        message_log = [
            self.system_message(
                f"You are a helpful data assistant. The user asked the question: '{question}'\n\n"
                f"The following is a pandas DataFrame with the results of the query: \n"
                f"{df.to_markdown()}\n\n"
            ),
            self.user_message(
                "Briefly summarize the data based on the question that was asked. "
                "Ensure that any ratios or percentages are formatted with the '%' symbol. "
                "Do not respond with any additional explanation beyond the summary. "
                "只需要提供最后的总结" +
                self._response_language()
            ),
        ]

        summary = self.submit_prompt(message_log, **kwargs)

        return summary

    # 这个函数是对Vanna的函数重写，作用是将输入的文本，制作成可以和ollama对话的格式
    def get_sql_prompt(
            self,
            initial_prompt: str,
            question: str,
            question_sql_list: list,
            ddl_list: list,
            doc_list: list,
            **kwargs,
    ):

        if initial_prompt is None:
            initial_prompt = f"You are a {self.dialect} expert. " + \
                             "Please help to generate a SQL query to answer the question. Your response should ONLY be based on the given context and follow the response guidelines and format instructions. "

        initial_prompt = self.add_ddl_to_prompt(
            initial_prompt, ddl_list, max_tokens=self.max_tokens
        )

        # if self.static_documentation != "":
        #     doc_list.append(self.static_documentation)

        initial_prompt = self.add_documentation_to_prompt(
            initial_prompt, doc_list, max_tokens=self.max_tokens
        )

        # initial_prompt += (
        #     "===Response Guidelines \n"
        #     "1. If the provided context is sufficient, please generate a valid SQL query without any explanations for the question. \n"
        #     "2. If the provided context is almost sufficient but requires knowledge of a specific string in a particular column, please generate an intermediate SQL query to find the distinct strings in that column. Prepend the query with a comment saying intermediate_sql \n"
        #     "3. If the provided context is insufficient, please explain why it can't be generated. \n"
        #     "4. Please use the most relevant table(s). \n"
        #     "5. If the question has been asked and answered before, please repeat the answer exactly as it was given before. \n"
        #     f"6. Ensure that the output SQL is {self.dialect}-compliant and executable, and free of syntax errors. \n"
        # )
        # "5. If the provided context is insufficient, please explain why it can't be generated. \n"
        initial_prompt += (

            "===Response Guidelines \n"
            "1. You are a database and knowledge assistant. You can answer data-related questions using the tables in the database, or answer general questions using your knowledge. Before answering, assess whether the question needs a database query. \n"
            "2. For conditions,  treat date fields as timestamps.\n"
            "3. If the provided context is sufficient, please generate a valid SQL query. \n"
            "4. If the provided context is almost sufficient but requires knowledge of a specific string in a particular column, please generate an intermediate SQL query to find the distinct strings in that column. Prepend the query with a comment saying intermediate_sql \n"
            "5. Please use the most relevant table(s). \n"
            "6. If the question has been asked and answered before, please repeat the answer exactly as it was given before. \n"
            f"7. Ensure that the output SQL is {self.dialect}-compliant and executable, and free of syntax errors. \n"
            "生成where的时候，主要不要搞错了字段的内容。\n"
            "SQL查询中将英文列名用AS转换为中文方便理解,中文记得加上引号\n"
            "一步一步的分析问题，生成sql来解决问题给出答案\n"
            "如果这个问题之前已经被问过并且回答过，请准确地重复之前给出的答案。\n"
            "生成 SQL 查询时，无论是做差值、汇总还是对比，请始终避免跨行计算。确保相关字段出现在同一行内进行运算，避免产生不准确的结果。\n"

        )  # "一定要记住，当需要计算环比和同比时，环比的group by里面不要包含月份，同比的group by里面不要包含年份\n"
        #   "生成where的时候，主要不要搞错了字段的内容,查询都要把日期和公司都查出来。\n"
        message_log = [self.system_message(initial_prompt)]
        # message_log=[]
        # for example in question_sql_list:
        #     if example is None:
        #         print("example is None")
        #     else:
        #         if example is not None and "question" in example and "sql" in example:
        #             message_log.append(self.user_message(example["question"]))
        #             message_log.append(self.assistant_message(example["sql"]))

        message_log.extend(kwargs.get('history', []))  # 提交历史记录

        message_log.append(self.user_message(question))  # 最新的提问

        return message_log

    # 这个函数是对Vanna里面的函数进行重写，作用是将ollama里的大模型的回答进行解析，提取回答里的sql，只获取文本最后的sql，因为那部分的sql才是完整的sql
    def extract_sql(self, llm_response):
        """
        Extracts the first SQL statement after the word 'select', ignoring case,
        matches until the first semicolon, three backticks, or the end of the string,
        and removes three backticks if they exist in the extracted string.

        Args:
        - llm_response (str): The string to search within for an SQL statement.

        Returns:
        - str: The first SQL statement found, with three backticks removed, or an empty string if no match is found.
        """
        # Remove ollama-generated extra characters
        llm_response = llm_response.replace("\\_", "_")
        llm_response = llm_response.replace("\\", "")

        # Regular expression to find ```sql' and capture until '```'
        # sql = re.search(r"```sql\n((.|\n)*?)(?=;|\[|```)", llm_response, re.DOTALL)
        # sql = re.findall(r"```sql\n((.|\n)*?)(?=;|\[|```)", llm_response, re.DOTALL)
        # Regular expression to find 'select, with (ignoring case) and capture until ';', [ (this happens in case of mistral) or end of string
        # select_with = re.search(r'(select|(with.*?as \())(.*?)(?=;|\[|```)',
        #                         llm_response,
        #                         re.IGNORECASE | re.DOTALL)
        # select_with=re.finditer(r'(select|(with.*?as \())(.*?)(?=;|\[|```)',
        #                         llm_response,
        #                         re.IGNORECASE | re.DOTALL)
        select_with = re.finditer(r'(?<=```sql\n)((.|\n)*?)(?=;|\[|```)',
                                  llm_response,
                                  re.IGNORECASE | re.DOTALL)
        # if sql:
        #     self.log(
        #         f"Output from LLM: {llm_response} \nExtracted SQL: {sql[-1][0]}")
        #     # return sql.group(1).replace("```", "")
        #     return sql[-1][0].replace("```", "")
        text = ""
        if select_with:

            # return select_with.group(0)
            # sql_query = ''.join(select_with[-1])
            for match in select_with:
                text = match.group(0)
            self.log(
                f"Output from LLM: {llm_response} \nExtracted SQL: {text}")
            return text
        else:
            return ""

    # 这个函数是自定义函数，这个用于将markdown格式的文本，制作成文档，包括设置文档标题等内容
    def create_word_from_markdown(self, markdown_text):
        """
        将 Markdown 文本内容解析为 Word 文档。
        :param markdown_text: 大模型生成的 Markdown 文本
        :return: Word 文档对象
        """
        document = Document()
        markdown_text = markdown_text.replace("markdown", "")
        markdown_text = markdown_text.replace("```", "")
        lines = markdown_text.split("\n")

        for line in lines:
            if line.startswith("# "):  # 一级标题
                document.add_heading(line[2:], level=1)
            elif line.startswith("## "):  # 二级标题
                document.add_heading(line[3:], level=2)
            elif line.startswith("### "):  # 二级标题
                document.add_heading(line[4:], level=3)
            elif line.startswith("#### "):  # 二级标题
                document.add_heading(line[5:], level=4)
            elif line.startswith("##### "):  # 二级标题
                document.add_heading(line[6:], level=5)
            elif line.startswith("- "):  # 列表项
                document.add_paragraph(line[2:], style="List Bullet")
            elif line.strip():  # 普通段落
                document.add_paragraph(line.strip())

        return document

    # 这个函数是自定义函数，这个和get_sql_prompt函数的作用是一样的，专门用于生成word这一功能，制作可以直接和ollama大模型对话的格式
    def generate_message_log(self, history_for_word):
        """
        根据历史对话记录生成 message_log。
        :param history_for_word: 用户选择的历史对话记录（列表）
        :return: 用于生成 Word 文档的 message_log
        """
        # 定义初始的系统消息，这里是对大模型角色的定义
        message_log = [
            self.system_message("你是一个熟练的财务分析师。根据以下对话历史记录，创建一个详细的 Word 报告，格式为 Markdown，报告包含以下部分：\n\n"
                                "1. **前言**：总结查询的背景和目的。\n"
                                "2. **查询内容**：详细列出用户的具体查询项。\n"
                                "3. **数据展示**：展示查询到的数据，包括时间变化趋势。\n"
                                "4. **深入分析**：对数据进行深入分析，例如资产、负债和所有者权益的变化。\n"
                                "5. **结论与建议**：总结发现并提出可操作的建议。")
        ]

        # 将历史记录格式化为用户和助手的对话，这个history_for_word是前端传过来的，用于制作生成文档的prompt
        for i in range(0, len(history_for_word), 2):  # 假设每两个元素是用户提问和大模型回答
            user_message = history_for_word[i]  # 用户提问
            assistant_message = history_for_word[i + 1] if i + 1 < len(history_for_word) else ""

            message_log.append(self.user_message(user_message))
            message_log.append(self.assistant_message(assistant_message))
        # 下面这个就是跟大模型提问的模板
        message_log.append(self.user_message(f"""
### 生成报告要求
你是一个财务分析专家，基于以下对话历史记录，生成一个详细的财务分析报告，格式为 Markdown，报告包含以下部分：

1. **前言**：总结查询的背景和目的。
2. **查询内容**：详细列出用户请求的查询项。
3. **数据展示**：展示查询到的数据，包含时间范围内的变化情况。
4. **深入分析**：对查询的数据进行详细分析.
5. **结论与建议**：总结数据的发现，并给出可操作的建议。

请严格按照这些部分进行总结，不要附加其他额外的解释或信息。输出格式应如下：

```
# 前言
[生成的前言内容]

# 查询内容
[生成的查询内容]

# 数据展示
[展示的数据内容]

# 深入分析
[生成的深入分析内容]

# 结论与建议
[生成的结论与建议内容]
```

请确保：
- 不要有额外的解释或多余的内容。
- 按照要求生成每一部分内容，不要漏掉任何部分。
- 输出必须符合 Markdown 格式，方便生成 Word 文档。
    """))
        return message_log

    # 这个是自定义函数，作用是专门为生成PPT这个功能制作和ollama大模型对话的格式
    def generate_message_log_ppt(self, history_for_word):
        """
        根据历史对话记录生成 message_log。
        :param history_for_word: 用户选择的历史对话记录（列表）
        :return: 用于生成 Word 文档的 message_log
        """
        # 定义初始的系统消息，这里是对大模型角色的定义
        message_log = [
            self.system_message("你是一个熟练的财务PPT生成专家。根据以下对话历史记录，创建一个详细的 PPT，格式为 json，PPT包含以下部分：\n\n"
                                "1. **前言**：总结查询的背景和目的。\n"
                                "2. **查询内容**：详细列出用户的具体查询项。\n"
                                "3. **数据展示**：展示查询到的数据，包括时间变化趋势。\n"
                                "4. **深入分析**：对数据进行深入分析，例如资产、负债和所有者权益的变化。\n"
                                "5. **结论与建议**：总结发现并提出可操作的建议。")
        ]

        # 将历史记录格式化为用户和助手的对话
        for i in range(0, len(history_for_word), 2):  # 假设每两个元素是用户提问和大模型回答
            user_message = history_for_word[i]  # 用户提问
            assistant_message = history_for_word[i + 1] if i + 1 < len(history_for_word) else ""

            message_log.append(self.user_message(user_message))
            message_log.append(self.assistant_message(assistant_message))
        message_log.append(self.user_message(
            """
### 生成PPT要求  
你是一位资深的财务 PPT 生成专家，基于以下对话历史记录，生成一个详细的财务分析 PPT，格式为 **JSON**。PPT 的内容需包括以下部分：  

1. **前言**  
   - 总结查询的背景和目的。  
2. **查询内容**  
   - 列出用户请求的具体查询项。  
3. **数据展示**  
   - 展示历史记录中查询到的数据，包含时间范围内的变化情况（如趋势、对比等）。  
   - 如果历史记录中有表格数据，请将其格式化为易于展示的方式（如简要摘要或数据点）。  
4. **深入分析**  
   - 基于展示的数据进行详细分析，包括发现的模式、潜在问题或其他相关洞察。  
5. **结论与建议**  
   - 总结主要发现，并提供具体、可操作的建议。  

### 输出要求  
生成的 JSON 必须符合以下格式：  

```json  
{  
    "title": "演示文稿标题",  
    "slides": [  
        {  
            "heading": "第一张幻灯片的标题",  
            "bullet_points": [  
                "第一项要点",  
                [  
                    "子要点 1",  
                    "子要点 2"  
                ],  
                "第二项要点"  
            ],  
            "key_message": ""  
        },  
        {  
            "heading": "第二张幻灯片的标题",  
            "bullet_points": [  
                "第一项要点",  
                "第二项要点",  
                "第三项要点"  
            ],  
            "key_message": "这张幻灯片传达的关键信息"  
        },  
        {  
            "heading": "描述逐步过程的幻灯片",  
            "bullet_points": [  
                ">> 过程的第一步（以特殊标记 >> 开头）",  
                ">> 第二步（以 >> 开头）",  
                ">> 第三步"  
            ],  
            "key_message": ""  
        },  
        {  
            "heading": "双栏布局的幻灯片（用于并排对比或对照两个相关的概念，例如优缺点、优势与风险、旧方法与现代方法等）",  
            "bullet_points": [  
                {  
                    "heading": "左栏标题",  
                    "bullet_points": [  
                        "第一项要点",  
                        "第二项要点",  
                        "第三项要点"  
                    ]  
                },  
                {  
                    "heading": "右栏标题",  
                    "bullet_points": [  
                        "第一项要点",  
                        "第二项要点",  
                        "第三项要点"  
                    ]  
                }  
            ],  
            "key_message": ""  
        }  
    ]  
}  
"""
        ))
        return message_log

    # 这个是对Vanna里面函数的重写，用于连接数据库，重点：已经适配了人大金仓数据库
    def connect_to_postgres(
            self,
            host: str = None,
            dbname: str = None,
            user: str = None,
            password: str = None,
            port: int = None,
            **kwargs
    ):

        """
        Connect to postgres using the psycopg2 connector. This is just a helper function to set [`vn.run_sql`][vanna.base.base.VannaBase.run_sql]
        **Example:**
        ```python
        vn.connect_to_postgres(
            host="myhost",
            dbname="mydatabase",
            user="myuser",
            password="mypassword",
            port=5432
        )
        ```
        Args:
            host (str): The postgres host.
            dbname (str): The postgres database name.
            user (str): The postgres user.
            password (str): The postgres password.
            port (int): The postgres Port.
        """

        try:
            import psycopg2
            import psycopg2.extras
        except ImportError:
            raise RuntimeError(
                "You need to install required dependencies to execute this method,"
                " run command: \npip install vanna[postgres]"
            )

        if not host:
            host = os.getenv("HOST")

        if not host:
            raise RuntimeError("Please set your postgres host")

        if not dbname:
            dbname = os.getenv("DATABASE")

        if not dbname:
            raise RuntimeError("Please set your postgres database")

        if not user:
            user = os.getenv("PG_USER")

        if not user:
            raise RuntimeError("Please set your postgres user")

        if not password:
            password = os.getenv("PASSWORD")

        if not password:
            raise RuntimeError("Please set your postgres password")

        if not port:
            port = os.getenv("PORT")

        if not port:
            raise RuntimeError("Please set your postgres port")

        conn = None

        try:
            conn = psycopg2.connect(
                host=host,
                dbname=dbname,
                user=user,
                password=password,
                port=port,
                **kwargs
            )
        except psycopg2.Error as e:
            raise RuntimeError(e)

        def connect_to_db():
            return psycopg2.connect(host=host, dbname=dbname,
                                    user=user, password=password, port=port, **kwargs)

        def run_sql_postgres(sql: str) -> Union[pd.DataFrame, None]:
            conn = None
            try:
                conn = connect_to_db()  # Initial connection attempt
                conn.set_client_encoding('UTF8')
                cs = conn.cursor()
                cs.execute(sql)
                results = cs.fetchall()

                # Create a pandas dataframe from the results
                df = pd.DataFrame(results, columns=[desc[0] for desc in cs.description])
                return df

            except psycopg2.InterfaceError as e:
                # Attempt to reconnect and retry the operation
                if conn:
                    conn.close()  # Ensure any existing connection is closed
                conn = connect_to_db()
                cs = conn.cursor()
                cs.execute(sql)
                results = cs.fetchall()

                # Create a pandas dataframe from the results
                df = pd.DataFrame(results, columns=[desc[0] for desc in cs.description])
                return df

            except psycopg2.Error as e:
                if conn:
                    conn.rollback()
                    raise RuntimeError(e)

            except Exception as e:
                conn.rollback()
                raise e

        self.dialect = "PostgreSQL"
        self.run_sql_is_set = True
        self.run_sql = run_sql_postgres

    # 这个是对Vanna里面函数的重写，用于让大模型生成plot代码，用于生成图表
    def generate_plotly_code(
            self, question: str = None, sql: str = None, df_metadata: str = None, **kwargs
    ) -> str:
        if question is not None:
            system_msg = f"The following is a pandas DataFrame that contains the results of the query that answers the question the user asked: '{question}'"
        else:
            system_msg = "The following is a pandas DataFrame "

        if sql is not None:
            system_msg += f"\n\nThe DataFrame was produced using this query: {sql}\n\n"

        system_msg += f"The following is information about the resulting pandas DataFrame 'df': \n{df_metadata}"

        message_log = [
            self.system_message(system_msg),
            self.user_message(
                "Can you generate the Python plotly code to chart the results of the dataframe? Assume the data is in a pandas dataframe called 'df'. If there is only one value in the dataframe, use an Indicator. Respond with only Python code and Chinese. Do not answer with any explanations -- just the code。"
            ),
        ]

        plotly_code = self.submit_prompt(message_log, kwargs=kwargs)

        return self._sanitize_plotly_code(self._extract_python_code(plotly_code))

    # 这个函数是对Vanna里面函数的重写，用于对大模型的回答里，提取python代码，这代码是生成图表的代码。同样只需要提取最终的代码，所以python_code[-1]是取的-1
    def _extract_python_code(self, markdown_string: str) -> str:
        # Regex pattern to match Python code blocks
        pattern = r"```[\w\s]*python\n([\s\S]*?)```|```([\s\S]*?)```"

        # Find all matches in the markdown string
        matches = re.findall(pattern, markdown_string, re.IGNORECASE)

        # Extract the Python code from the matches
        python_code = []
        for match in matches:
            python = match[0] if match[0] else match[1]
            python_code.append(python.strip())

        if len(python_code) == 0:
            return markdown_string

        return python_code[-1]


# 这个是创建嵌入模型函数的对象
embedding_function = OllamaEmbeddingFunction(model_name="bge-m3:latest")
# "embedding_function":embedding_function,
vn = MyVanna(config={"embedding_function": embedding_function, 'model': 'qwq:32b', "n_results_sql": 4,
                     "n_results_documentation": 1, "n_results_ddl": 1, 'options': {"num_ctx": 6000},
                     "max_tokens": 32768, "language": "chinese", "initial_prompt":
                         f"你是一个数据库和知识系统的助手，你可以根据数据库中的表来回答数据相关问题，也可以通过你的知识回答一般性问题。\n"
                         f"请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL；\n"
                         f"如果不需要，请直接回答\n"
                         f"日期字段的属性是timestamp,\n"
                         "生成where的时候，主要不要搞错了字段的内容\n"
                         "SQL查询中将英文列名用AS转换为中文方便理解,中文记得加上引号\n"
                         "一步一步的分析问题，生成sql来解决问题给出答案\n"
                         "生成 SQL 查询时，无论是做差值、汇总还是对比，请始终避免跨行计算。确保相关字段出现在同一行内进行运算，避免产生不准确的结果。\n"})
# 连接数据库，数据库部署在阿里云上
vn.connect_to_postgres(host='223.4.251.109', dbname='kingbase', user='kingbase', password='123456', port=54321)
# 这些是用于程序运行中的变量
history = []
doc_list_tmp = []
ddl_list_tmp = []
data_rag = {}
# region 对数据表进行加载
# 下面是对每张数据表进行背景知识的构建，对每张表进行详细的解释，然后存如data_rag字典里去，之后会根据chromdb的检索结果，来从这里面取特定的一张表，输入到大模型对话里面去
# 后面是对8张表进行一样的处理
# 表cwgk_report_vb_cwys031
data_rag['cwgk_report_vb_cwys031'] = []
data_rag['cwgk_report_vb_cwys031'].append(
    '''-- public.cwgk_report_vb_cwys031 definition
    
    -- Drop table
    
    -- DROP TABLE public.cwgk_report_vb_cwys031;
    
    CREATE TABLE public.cwgk_report_vb_cwys031 (
        organ_id varchar(255) NULL, -- 组织id
        organname varchar(255) NULL, -- 编制单位
        item_id float8 NULL, -- 科目id
        item_name varchar(255) NULL, -- 科目名
        rq timestamp NULL, -- 交易日期
        update_time timestamp NULL, -- 上传日期
        jldw varchar(255) NULL, -- 单位
        ysbkhb float8 NULL, -- 预算板块合并
        yshdhb float8 NULL, -- 预算火电合并
        ysrm float8 NULL, -- 预算燃煤
        ysrq float8 NULL, -- 预算燃气
        ysxny float8 NULL, -- 预算新能源
        xzbkhb float8 NULL, -- 修正板块合并
        xzhdhb float8 NULL, -- 修正火电合并
        xzrm float8 NULL, -- 修正燃煤
        xzrq float8 NULL, -- 修正燃气
        xzxny float8 NULL, -- 修正新能源
        sjbkhb float8 NULL, -- 实际板块合并
        sjhdhb float8 NULL, -- 实际火电合并
        sjrm float8 NULL, -- 实际燃煤
        sjrq float8 NULL, -- 实际燃气
        sjxny float8 NULL, -- 实际新能源
        kongbai float8 NULL -- 空白
    );
    
    -- Column comments
    
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.organ_id IS '组织id';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.organname IS '编制单位';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.item_id IS '科目id';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.item_name IS '科目名';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.rq IS '交易日期';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.update_time IS '上传日期';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.jldw IS '单位';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.ysbkhb IS '预算板块合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.yshdhb IS '预算火电合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.ysrm IS '预算燃煤';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.ysrq IS '预算燃气';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.ysxny IS '预算新能源';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.xzbkhb IS '修正板块合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.xzhdhb IS '修正火电合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.xzrm IS '修正燃煤';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.xzrq IS '修正燃气';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.xzxny IS '修正新能源';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.sjbkhb IS '实际板块合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.sjhdhb IS '实际火电合并';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.sjrm IS '实际燃煤';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.sjrq IS '实际燃气';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.sjxny IS '实际新能源';
    COMMENT ON COLUMN public.cwgk_report_vb_cwys031.kongbai IS '空白';
    '''
)
data_rag['cwgk_report_vb_cwys031'].append(
    """cwgk_report_vb_cwys031 表是一个月度预算和财务科目记录表，记录每个月各个公司在不同科目下的预算、修正和实际金额。该表的主要列及其含义如下：
    1. organ_id: 组织 ID，用于标识所属的公司或组织。
    2. organname: 编制单位，标识公司或机构的名称。
    3. item_id: 科目 ID，表示财务科目的唯一标识符。
    4. item_name: 科目名，描述该记录所涉及的财务科目，例如“燃煤”、“燃气”等。
    5. rq: 交易日期，记录该数据的交易时间。
    6. update_time: 上传日期，记录该数据的上传时间。
    7. jldw: 记录的单位，表示财务数据使用的计量单位。
    8. ysbkhb: 预算板块合并，记录公司的预算板块合并情况。
    9. yshdhb: 预算火电合并，记录公司的火电预算合并情况。
    10. ysrm: 预算燃煤，记录公司在燃煤方面的预算金额。
    11. ysrq: 预算燃气，记录公司在燃气方面的预算金额。
    12. ysxny: 预算新能源，记录公司在新能源方面的预算金额。
    13. xzbkhb: 修正板块合并，记录预算修正后的板块合并数据。
    14. xzhdhb: 修正火电合并，记录修正后的火电合并预算数据。
    15. xzrm: 修正燃煤，记录修正后的燃煤预算数据。
    16. xzrq: 修正燃气，记录修正后的燃气预算数据。
    17. xzxny: 修正新能源，记录修正后的新能源预算数据。
    18. sjbkhb: 实际板块合并，记录实际的板块合并金额。
    19. sjhdhb: 实际火电合并，记录实际的火电合并金额。
    20. sjrm: 实际燃煤，记录公司在燃煤方面的实际支出。
    21. sjrq: 实际燃气，记录公司在燃气方面的实际支出。
    22. sjxny: 实际新能源，记录公司在新能源方面的实际支出。
    23. kongbai: 空白字段，用于补位或其他扩展。
    
    这些列中，'item_name和organname' 列是关键字段之一，item_name包含了多个固定的财务科目名称，如燃煤、燃气等。
    organname列包含了多个固定的公司名称。当用户询问与特定科目相关的预算或支出时，模型可以根据 'item_name'和organname的值来查询相应的数据。\n"""
)
# 这里是先查询cwgk_report_vb_cwys031表，查询两个字段的内容，item_names是指标内容，organname是公司内容，这些内容大模型必须要知道，才可以生成正确的sql
item_names = vn.run_sql("SELECT DISTINCT item_name FROM cwgk_report_vb_cwys031 WHERE item_name IS NOT NULL ")
organname = vn.run_sql("SELECT DISTINCT organname FROM cwgk_report_vb_cwys031 WHERE organname IS NOT NULL ")

item_names_list = [item for item in item_names['item_name'].tolist() if item is not None]
organname_list = [item for item in organname['organname'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
item_names_str = ','.join(item_names_list)
organname_str = ','.join(organname_list)
data_rag['cwgk_report_vb_cwys031'].append(
    """cwgk_report_vb_cwys031表
       """ + f"""**关键字段**:
1. **item_name** (科目名):
   - **固定取值列表**: 
     - [item_names_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organname** (组织名称):
   - **固定取值列表**:
     - [{organname_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **item_name** 和 **organname** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
# 这个是将item_names_list内容存到本地数据chromadb中去
vn.train_yu(text=item_names_list, name='cwgk_report_vb_cwys031')

# #表sap_yzb0003
data_rag['sap_yzb0003'] = []
data_rag['sap_yzb0003'].append(
    """
    -- public.sap_yzb0003 definition
    
    -- Drop table
    
    -- DROP TABLE public.sap_yzb0003;
    
    CREATE TABLE public.sap_yzb0003 (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 交易日期
        proj_name varchar(255) NULL, -- 科目名称
        seq_no float8 NULL, -- 行次
        curr_y_amt float8 NULL, -- 本年金额
        last_y_amt float8 NULL, -- 上年金额
        uploadtime timestamp NULL -- 上传日期
    );
    
    -- Column comments
    
    COMMENT ON COLUMN public.sap_yzb0003.organ_code IS '组织代码';
    COMMENT ON COLUMN public.sap_yzb0003.organ_name IS '组织名称';
    COMMENT ON COLUMN public.sap_yzb0003.tdate IS '录入日期';
    COMMENT ON COLUMN public.sap_yzb0003.proj_name IS '科目名称';
    COMMENT ON COLUMN public.sap_yzb0003.seq_no IS '行次';
    COMMENT ON COLUMN public.sap_yzb0003.curr_y_amt IS '本年金额';
    COMMENT ON COLUMN public.sap_yzb0003.last_y_amt IS '上年金额';
    COMMENT ON COLUMN public.sap_yzb0003.uploadtime IS '更新日期';
    """
)
data_rag['sap_yzb0003'].append(
    """sap_yzb0003 是一张现金流量表，记录每个公司在不同款项下的本年和上年金额。该表的设计目的是用于追踪各公司财务科目在不同年度的现金流变化。表的主要列及其含义如下：
    
    1.organ_code (组织代码)：表示每个公司的唯一标识符，通常是字母和数字的组合，用于区分不同的公司或组织。
    
    2.organ_name (组织名称)：公司或组织的名称，用于标识现金流属于哪个公司。
    
    3.tdate (交易日期)：表示该笔交易的日期，记录现金流的实际发生时间。
    
    4.proj_name (科目名称)：表示与该笔现金流相关的财务科目名称。用于描述该现金流所属的具体财务项目或类别。
    
    5.seq_no (行次)：记录该笔交易在表中的行次，用于对表中的数据进行排序或快速定位特定条目。
    
    6.curr_y_amt (本年金额)：记录该科目下当前年度的现金流金额。用于跟踪本年度的现金流动。
    
    7.last_y_amt (上年金额)：记录该科目下上一年度的现金流金额，用于与当前年度进行对比分析。
    
    8.uploadtime (上传日期)：表示数据的最后上传日期或更新时间，确保数据的时间戳完整，便于数据版本控制。\n"""
)
proj_name_2 = vn.run_sql("SELECT DISTINCT proj_name FROM sap_yzb0003 WHERE proj_name IS NOT NULL ")
organ_name_2 = vn.run_sql("SELECT DISTINCT organ_name FROM sap_yzb0003 WHERE organ_name IS NOT NULL ")

proj_name_2_list = [item for item in proj_name_2['proj_name'].tolist() if item is not None]
organ_name_2_list = [item for item in organ_name_2['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
proj_name_2_str = ','.join(proj_name_2_list)
organ_name_2_str = ','.join((organ_name_2_list))

data_rag['sap_yzb0003'].append(
    """
    sap_yzb0003表
    """ +
    f"""**关键字段**:
1. **proj_name** (科目名):
   - **固定取值列表**: 
     - [proj_name_2_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_2_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **proj_name** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=proj_name_2_list, name='sap_yzb0003')

data_rag['twb_ygl0078'] = []
data_rag['twb_ygl0078'].append(
    """
    -- public.twb_ygl0078 definition
    
    -- Drop table
    
    -- DROP TABLE public.twb_ygl0078;
    
    CREATE TABLE public.twb_ygl0078 (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 交易日期
        ind_cate_name varchar(255) NULL, -- 指标类别名称
        mat_ind_name varchar(255) NULL, -- 指标名称
        unit_name varchar(255) NULL, -- 单位
        curr_y_budget float8 NULL, -- 年度预算
        last_m_pe_num float8 NULL, -- 上月
        curr_m_pe_num float8 NULL, -- 本月
        last_ym_pe_num float8 NULL, -- 上年同期
        curr_y_pe_num float8 NULL, -- 本年累计
        uploadtime timestamp NULL -- 上传日期
    );
    
    -- Column comments
    
    COMMENT ON COLUMN public.twb_ygl0078.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twb_ygl0078.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twb_ygl0078.tdate IS '交易日期';
    COMMENT ON COLUMN public.twb_ygl0078.ind_cate_name IS '指标类别名称';
    COMMENT ON COLUMN public.twb_ygl0078.mat_ind_name IS '指标名称';
    COMMENT ON COLUMN public.twb_ygl0078.unit_name IS '单位';
    COMMENT ON COLUMN public.twb_ygl0078.curr_y_budget IS '年度预算';
    COMMENT ON COLUMN public.twb_ygl0078.last_m_pe_num IS '上月';
    COMMENT ON COLUMN public.twb_ygl0078.curr_m_pe_num IS '本月';
    COMMENT ON COLUMN public.twb_ygl0078.last_ym_pe_num IS '上年同期';
    COMMENT ON COLUMN public.twb_ygl0078.curr_y_pe_num IS '本年累计';
    COMMENT ON COLUMN public.twb_ygl0078.uploadtime IS '上传日期';
    """
)
data_rag['twb_ygl0078'].append(
    """twb_ygl0078是一张集团汇总指标的数据表。该表专注于汇总各个组织的关键财务和经营指标，提供年度预算及各类时间维度下的数据表现，如本月、上月、上年同期和本年累计的数据。各列详细说明如下：
    
    1.organ_code (varchar): 组织代码，用于唯一标识集团下的各个子公司或组织。
    2.organ_name (varchar): 组织名称，记录子公司或组织的名称。
    3.tdate (timestamp): 交易日期，记录数据所属的具体日期。
    4.ind_cate_name (varchar): 指标类别名称，标识数据所属的类别（如财务、销售等类别）。
    5.mat_ind_name (varchar): 指标名称，描述具体的指标（如年度预算、本月支出等）。
    6.unit_name (varchar): 单位，记录指标的计量单位（如万元、件等）。
    7.curr_y_budget (float8): 年度预算，记录该组织或子公司在当前年度的预算数额。
    8.last_m_pe_num (float8): 上月，记录该组织在上月的指标数据。
    9.curr_m_pe_num (float8): 本月，记录该组织在本月的指标数据。
    10.last_ym_pe_num (float8): 上年同期，记录该组织在上年同期的指标数据。
    11.curr_y_pe_num (float8): 本年累计，记录从1月到本月的累计值。
    12.uploadtime (timestamp): 上传日期，表示该数据的最后更新时间。\n"""
)
#
#
mat_ind_name_3 = vn.run_sql("SELECT DISTINCT mat_ind_name FROM twb_ygl0078 WHERE mat_ind_name IS NOT NULL ")
organ_name_3 = vn.run_sql("SELECT DISTINCT organ_name FROM twb_ygl0078 WHERE organ_name IS NOT NULL ")
# ind_cate_name_3 = vn.run_sql("SELECT DISTINCT ind_cate_name FROM twb_ygl0078 WHERE ind_cate_name IS NOT NULL ")
mat_and_ind_name_3 = vn.run_sql("SELECT DISTINCT mat_ind_name,ind_cate_name FROM twb_ygl0078")
mat_and_ind_name_3_table = mat_and_ind_name_3.to_markdown(index=False)
# print(mat_and_ind_name_3_table)
mat_ind_name_3_list = [item for item in mat_ind_name_3['mat_ind_name'].tolist() if item is not None]
organ_name_3_list = [item for item in organ_name_3['organ_name'].tolist() if item is not None]
# ind_cate_name_3_list=[item for item in ind_cate_name_3['ind_cate_name'].tolist() if item is not None]

# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
# mat_ind_name_3_str = ', '.join(mat_ind_name_3_list)
organ_name_3_str = ', '.join(organ_name_3_list)
# ind_cate_name_3_str=', '.join(ind_cate_name_3_list)
data_rag['twb_ygl0078'].append(
    """
    twb_ygl0078表
    """ +
    f"""**关键字段**:
1. **organ_name** (组织名):
   - **固定取值列表**: 
     - [{organ_name_3_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
**注意**:
- **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。\n"""
    "以下是mat_ind_name和ind_cate_name字段的内容：\n\n"
    f"{mat_and_ind_name_3_table}\n\n"
    "请根据这些数据生成 SQL 查询，并确保 WHERE 子句中的值和列搭配一致。"
)
vn.train_yu(text=mat_ind_name_3_list, name='twb_ygl0078')

#
# #表twb_yzb0102
data_rag['twb_yzb0102'] = []
data_rag['twb_yzb0102'].append(
    """
    -- public.twb_yzb0102 definition
    
    -- Drop table
    
    -- DROP TABLE public.twb_yzb0102;
    
    CREATE TABLE public.twb_yzb0102 (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 交易日期
        asset_desc varchar(255) NULL, -- 科目名称
        item_id float8 NULL, -- id号
        closing_balance float8 NULL, -- 期末金额
        last_ym_pe_num float8 NULL, -- 上年同期数
        year_begin_bal float8 NULL, -- 年初金额
        uploadtime timestamp NULL -- 上传日期
    );
    COMMENT ON TABLE public.twb_yzb0102 IS '资产负债表';
    
    -- Column comments
    
    COMMENT ON COLUMN public.twb_yzb0102.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twb_yzb0102.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twb_yzb0102.tdate IS '交易日期';
    COMMENT ON COLUMN public.twb_yzb0102.asset_desc IS '科目名称';
    COMMENT ON COLUMN public.twb_yzb0102.item_id IS 'id号';
    COMMENT ON COLUMN public.twb_yzb0102.closing_balance IS '期末金额';
    COMMENT ON COLUMN public.twb_yzb0102.last_ym_pe_num IS '上年同期数';
    COMMENT ON COLUMN public.twb_yzb0102.year_begin_bal IS '年初金额';
    COMMENT ON COLUMN public.twb_yzb0102.uploadtime IS '上传日期';
    """
)
data_rag['twb_yzb0102'].append(
    """twb_yzb0102 表是一个资产负债表，专门记录每个公司或组织的资产项目及其相关金额。该表的主要列及其含义如下：
    
    1.organ_code: 组织代码，用于标识所属的公司或组织。
    2.organ_name: 组织名称，表示记录所属的公司或机构的名称。
    3.tdate: 交易日期，记录该数据的录入或生效时间。
    4.asset_desc: 科目名称，描述该记录所涉及的资产科目，例如“固定资产”、“流动资产”等。
    5.item_id: ID号，表示资产项目的唯一标识符。
    6.closing_balance: 期末金额，表示截至期末该科目下的资产金额。
    7.last_ym_pe_num: 上年同期数，表示去年同期该资产项目的金额。
    8.year_begin_bal: 年初金额，表示该科目在年初的资产金额。
    9.uploadtime: 上传日期，记录该数据的最后更新时间。\n"""
)
#
asset_desc_4 = vn.run_sql("SELECT DISTINCT asset_desc FROM twb_yzb0102 WHERE asset_desc IS NOT NULL ")
organ_name_4 = vn.run_sql("SELECT DISTINCT organ_name FROM twb_yzb0102 WHERE organ_name IS NOT NULL ")

asset_desc_4_list = [item for item in asset_desc_4['asset_desc'].tolist() if item is not None]
organ_name_4_list = [item for item in organ_name_4['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
asset_desc_4_str = ', '.join(asset_desc_4_list)
organ_name_4_str = ', '.join((organ_name_4_list))
data_rag['twb_yzb0102'].append(
    """
    twb_yzb0102表
    """ +
    f"""**关键字段**:
1. **asset_desc** (科目名):
   - **固定取值列表**: 
     - [asset_desc_4_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_4_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **asset_desc** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=asset_desc_4_list, name='twb_yzb0102')

# #twc_km_bal_m
data_rag['twc_km_bal_m'] = []
data_rag['twc_km_bal_m'].append(
    """
    -- public.twc_km_bal_m definition
    
    -- Drop table
    
    -- DROP TABLE public.twc_km_bal_m;
    
    CREATE TABLE public.twc_km_bal_m (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 录入日期
        data_code varchar(255) NULL, -- 数据代码
        item_code float8 NULL, -- id代码
        item_name varchar(255) NULL, -- 科目名称
        func_arrange float8 NULL, -- 功能安排
        beg_money_direct varchar(255) NULL, -- 期初余额方向
        beg_money_bal float8 NULL, -- 期初余额
        curr_m_j float8 NULL, -- 本月应交
        curr_m_d float8 NULL, -- 本月重分类
        acc_j float8 NULL, -- 本年应交
        acc_d float8 NULL, -- 本年重分类
        end_direction float8 NULL, -- 期末余额方向
        end_bal float8 NULL, -- 期末余额
        profit_center varchar(255) NULL, -- 利润部门
        profit_center_desc varchar(255) NULL, -- 利润部门描述
        money_id varchar(255) NULL, -- 货币单位
        beg_y_money_direct varchar(255) NULL, -- 年初余额重分类转成
        beg_y_foreign_bal float8 NULL, -- 年初外币余额
        beg_y_money_bal float8 NULL, -- 年初余额
        beg_foreign_bal float8 NULL, -- 期初外币余额
        curr_m_foreign_j float8 NULL, -- 本月应交外币金额
        curr_m_foreign_d float8 NULL, -- 本月重分类外币金额
        curr_m_pe_num_f float8 NULL, -- 本月外币已交金额
        curr_m_pe_num float8 NULL, -- 本月已交
        acc_foreign_j float8 NULL, -- 本年应交外币金额
        acc_foreign_d float8 NULL, -- 本年重分类外币金额
        acc_pe_foreign float8 NULL, -- 本年外币已交金额
        acc_pe_num float8 NULL, -- 本年已交
        end_foreign_bal float8 NULL, -- 期末外币余额
        uploadtime timestamp NULL -- 上传日期
    );
    COMMENT ON TABLE public.twc_km_bal_m IS '科目余额表';
    
    -- Column comments
    
    COMMENT ON COLUMN public.twc_km_bal_m.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twc_km_bal_m.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twc_km_bal_m.tdate IS '录入日期';
    COMMENT ON COLUMN public.twc_km_bal_m.data_code IS '数据代码';
    COMMENT ON COLUMN public.twc_km_bal_m.item_code IS 'id代码';
    COMMENT ON COLUMN public.twc_km_bal_m.item_name IS '科目名称';
    COMMENT ON COLUMN public.twc_km_bal_m.func_arrange IS '功能安排';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_money_direct IS '期初余额方向';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_money_bal IS '期初余额';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_j IS '本月应交';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_d IS '本月重分类';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_j IS '本年应交';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_d IS '本年重分类';
    COMMENT ON COLUMN public.twc_km_bal_m.end_direction IS '期末余额方向';
    COMMENT ON COLUMN public.twc_km_bal_m.end_bal IS '期末余额';
    COMMENT ON COLUMN public.twc_km_bal_m.profit_center IS '利润部门';
    COMMENT ON COLUMN public.twc_km_bal_m.profit_center_desc IS '利润部门描述';
    COMMENT ON COLUMN public.twc_km_bal_m.money_id IS '货币单位';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_y_money_direct IS '年初余额重分类转成';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_y_foreign_bal IS '年初外币余额';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_y_money_bal IS '年初余额';
    COMMENT ON COLUMN public.twc_km_bal_m.beg_foreign_bal IS '期初外币余额';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_foreign_j IS '本月应交外币金额';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_foreign_d IS '本月重分类外币金额';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_pe_num_f IS '本月外币已交金额';
    COMMENT ON COLUMN public.twc_km_bal_m.curr_m_pe_num IS '本月已交';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_foreign_j IS '本年应交外币金额';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_foreign_d IS '本年重分类外币金额';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_pe_foreign IS '本年外币已交金额';
    COMMENT ON COLUMN public.twc_km_bal_m.acc_pe_num IS '本年已交';
    COMMENT ON COLUMN public.twc_km_bal_m.end_foreign_bal IS '期末外币余额';
    COMMENT ON COLUMN public.twc_km_bal_m.uploadtime IS '上传日期';
    """
)
data_rag['twc_km_bal_m'].append(
    """twc_km_bal_m 是一个用于记录各个公司的科目余额的表，主要包含期初余额、本月应交、本月重分类、本年应交、本年重分类和期末余额等信息。该表的主要列及其含义如下：
    
    1.organ_code: 组织代码，用于标识所属的公司或组织。
    2.organ_name: 组织名称，标识公司或机构的名称。
    3.tdate: 录入日期，记录数据的录入时间。
    4.data_code: 数据代码，用于标识具体数据的代码。
    5.item_code: 科目 ID，表示财务科目的唯一标识符。
    6.item_name: 科目名称，描述财务科目，例如“应交税费”、“应收账款”等。
    7.func_arrange: 功能安排，记录相关的功能信息。
    8.beg_money_direct: 期初余额方向，表示期初余额的正负方向。
    9.beg_money_bal: 期初余额，表示在本期初始时的余额金额。
    10.curr_m_j: 本月应交，记录本月的应交金额。
    11.curr_m_d: 本月重分类，记录本月的重分类金额。
    12.acc_j: 本年应交，记录本年度累计的应交金额。
    13.acc_d: 本年重分类，记录本年度累计的重分类金额。
    14.end_direction: 期末余额方向，表示期末余额的正负方向。
    15.end_bal: 期末余额，记录在本期结束时的余额金额。
    16.profit_center: 利润部门，表示相关利润中心的代码。
    17.profit_center_desc: 利润部门描述，记录利润中心的详细说明。
    18.money_id: 货币单位，用于表示使用的货币类型。
    19.beg_y_money_direct: 年初余额重分类转成，记录年初余额的重分类方向。
    20.beg_y_foreign_bal: 年初外币余额，记录以外币表示的年初余额。
    21.beg_y_money_bal: 年初余额，记录年初时的余额金额。
    22.beg_foreign_bal: 期初外币余额，记录以外币表示的期初余额。
    23.curr_m_foreign_j: 本月应交外币金额，记录本月应交的外币金额。
    24.curr_m_foreign_d: 本月重分类外币金额，记录本月重分类的外币金额。
    25.curr_m_pe_num_f: 本月外币已交金额，记录本月已交的外币金额。
    26.curr_m_pe_num: 本月已交，记录本月已交金额。
    27.acc_foreign_j: 本年应交外币金额，记录本年应交的外币金额。
    28.acc_foreign_d: 本年重分类外币金额，记录本年重分类的外币金额。
    29.acc_pe_foreign: 本年外币已交金额，记录本年已交的外币金额。
    30.acc_pe_num: 本年已交，记录本年已交金额。
    31.end_foreign_bal: 期末外币余额，记录以外币表示的期末余额。
    32.uploadtime: 上传日期，记录数据的最后上传时间。
    该表中的关键字段包括 organ_name（组织名称）、item_name（科目名称）、以及余额相关的列如 beg_money_bal（期初余额）和 end_bal（期末余额），用于跟踪不同公司的财务科目余额及其变化。\n"""
)
item_name_5 = vn.run_sql("SELECT DISTINCT item_name FROM twc_km_bal_m WHERE item_name IS NOT NULL ")
organ_name_5 = vn.run_sql("SELECT DISTINCT organ_name FROM twc_km_bal_m WHERE organ_name IS NOT NULL ")

item_name_5_list = [item for item in item_name_5['item_name'].tolist() if item is not None]
organ_name_5_list = [item for item in organ_name_5['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
item_name_5_str = ', '.join(item_name_5_list)
organ_name_5_str = ', '.join((organ_name_5_list))
data_rag['twc_km_bal_m'].append(
    """
    twc_km_bal_m表
    """ +
    f"""**关键字段**:
1. **item_name** (科目名):
   - **固定取值列表**: 
     - [item_name_5_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_5_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **item_name** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=item_name_5_list, name='twc_km_bal_m')

# #twc_profit_state_m
data_rag['twc_profit_state_m'] = []
data_rag['twc_profit_state_m'].append(
    """
    -- public.twc_profit_state_m definition
    
    -- Drop table
    
    -- DROP TABLE public.twc_profit_state_m;
    
    CREATE TABLE public.twc_profit_state_m (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 交易日期
        data_code varchar(255) NULL, -- 数据代码
        proj_name varchar(255) NULL, -- 科目名称
        item_id int4 NULL, -- id号
        curr_m_pe_num float8 NULL, -- 本月数
        curr_y_pe_num float8 NULL, -- 本年累计数
        last_ym_pe_num float8 NULL, -- 上年同期数
        uploadtime timestamp NULL -- 上传日期
    );
    COMMENT ON TABLE public.twc_profit_state_m IS '利润表';
    
    -- Column comments
    
    COMMENT ON COLUMN public.twc_profit_state_m.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twc_profit_state_m.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twc_profit_state_m.tdate IS '交易日期';
    COMMENT ON COLUMN public.twc_profit_state_m.data_code IS '数据代码';
    COMMENT ON COLUMN public.twc_profit_state_m.proj_name IS '科目名称';
    COMMENT ON COLUMN public.twc_profit_state_m.item_id IS 'id号';
    COMMENT ON COLUMN public.twc_profit_state_m.curr_m_pe_num IS '本月数';
    COMMENT ON COLUMN public.twc_profit_state_m.curr_y_pe_num IS '本年累计数';
    COMMENT ON COLUMN public.twc_profit_state_m.last_ym_pe_num IS '上年同期数';
    COMMENT ON COLUMN public.twc_profit_state_m.uploadtime IS '上传日期';
    """
)
data_rag['twc_profit_state_m'].append(
    """twc_profit_state_m是一张利润表，该表用于存储公司每月和年度累计的利润信息，帮助分析当前月和上年同期的对比。通过该表，用户可以查询特定公司在某一时间段内的利润数据，包括按月和按年累计的利润情况。
    该表的主要列及其含义如下：
    1.organ_code (varchar(255)): 组织代码，用于标识所属的公司或组织。
    2.organ_name (varchar(255)): 组织名称，标识公司或机构的名称。
    3.tdate (timestamp): 交易日期，表示该数据的记录时间。
    4.data_code (varchar(255)): 数据代码，标识该条利润记录的唯一数据代码。
    5.proj_name (varchar(255)): 科目名称，描述该记录所涉及的利润科目。
    6.item_id (int4): ID号，表示利润表项目的唯一标识符。
    7.curr_m_pe_num (float8): 本月数，表示当前月的利润金额。
    8.curr_y_pe_num (float8): 本年累计数，表示当前年份的累计利润金额。
    9.last_ym_pe_num (float8): 上年同期数，表示上一年度相同时间的利润金额。
    10.uploadtime (timestamp): 上传日期，记录该数据的最后上传时间。\n"""
)
#
proj_name_6 = vn.run_sql("SELECT DISTINCT proj_name FROM twc_profit_state_m WHERE proj_name IS NOT NULL ")
organ_name_6 = vn.run_sql("SELECT DISTINCT organ_name FROM twc_profit_state_m WHERE organ_name IS NOT NULL ")

proj_name_6_list = [item for item in proj_name_6['proj_name'].tolist() if item is not None]
organ_name_6_list = [item for item in organ_name_6['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
proj_name_6_str = ', '.join(proj_name_6_list)
organ_name_6_str = ', '.join((organ_name_6_list))

data_rag['twc_profit_state_m'].append(
    """
    twc_profit_state_m表
    """ +
    f"""**关键字段**:
1. **proj_name** (科目名):
   - **固定取值列表**: 
     - [proj_name_6_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_6_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **proj_name** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=proj_name_6_list, name='twc_profit_state_m')

# #twc_thermal_power_operate_m
data_rag['twc_thermal_power_operate_m'] = []
data_rag['twc_thermal_power_operate_m'].append(
    """
    -- public.twc_thermal_power_operate_m definition
    
    -- Drop table
    
    -- DROP TABLE public.twc_thermal_power_operate_m;
    
    CREATE TABLE public.twc_thermal_power_operate_m (
        organ_code varchar(255) NULL, -- 组织代码
        organ_name varchar(255) NULL, -- 组织名称
        tdate timestamp NULL, -- 录入日期
        data_code varchar(255) NULL, -- 数据代码
        proj_name varchar(255) NULL, -- 科目名称
        item_id float8 NULL, -- id号
        unit varchar(255) NULL, -- 单位
        curr_y_budget float8 NULL, -- 本年预算总金额
        last_m_pe_num float8 NULL, -- 上月完成总金额
        curr_m_pe_num float8 NULL, -- 本月完成总金额
        last_ym_pe_num float8 NULL, -- 上年同期总金额
        curr_y_pe_num float8 NULL, -- 本年累计总金额
        curr_y_same_pe_num float8 NULL, -- 同比增减额总金额
        curr_y_same_pe_factor float8 NULL, -- 同比增减利因素的影响原因
        curr_y_budegt_same_num float8 NULL, -- 本年预算同比增减单位金额
        curr_y_budegt_same_factor float8 NULL, -- 本年预算同比增减的原因
        uploadtime timestamp NULL, -- 上传日期
        curr_y_budget_unit float8 NULL, -- 本年预算单位金额
        last_m_pe_num_unit float8 NULL, -- 上月完成单位金额
        curr_m_pe_num_unit float8 NULL, -- 本月完成单位金额
        curr_y_pe_num_unit float8 NULL, -- 本年累计单位金额
        last_ym_pe_num_unit float8 NULL, -- 上年同期单位金额
        curr_y_same_pe_num_unit float8 NULL, -- 同比增减额单位金额
        curr_y_same_pe_factor_unit float8 NULL, -- 同比增减利因素的影响金额
        cause_factor float8 NULL -- 同比增减利因素的影响原因
    );
    COMMENT ON TABLE public.twc_thermal_power_operate_m IS '火力发电经营情况表';
    
    -- Column comments
    
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.tdate IS '录入日期';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.data_code IS '数据代码';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.proj_name IS '科目名称';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.item_id IS 'id号';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.unit IS '单位';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_budget IS '本年预算总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.last_m_pe_num IS '上月完成总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_m_pe_num IS '本月完成总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.last_ym_pe_num IS '上年同期总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_pe_num IS '本年累计总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_same_pe_num IS '同比增减额总金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_same_pe_factor IS '同比增减利因素的影响原因';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_budegt_same_num IS '本年预算同比增减单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_budegt_same_factor IS '本年预算同比增减的原因';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.uploadtime IS '上传日期';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_budget_unit IS '本年预算单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.last_m_pe_num_unit IS '上月完成单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_m_pe_num_unit IS '本月完成单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_pe_num_unit IS '本年累计单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.last_ym_pe_num_unit IS '上年同期单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_same_pe_num_unit IS '同比增减额单位金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.curr_y_same_pe_factor_unit IS '同比增减利因素的影响金额';
    COMMENT ON COLUMN public.twc_thermal_power_operate_m.cause_factor IS '同比增减利因素的影响原因';
    """
)
data_rag['twc_thermal_power_operate_m'].append(
    """
    twc_thermal_power_operate_m是火力发电经营情况表。
    该表用于记录火力发电公司各项经营指标的预算、完成情况以及同比变化。通过该表，用户可以分析公司经营的进展、预算的执行情况，并追踪同比增减的因素及其影响。
    该表的主要列及其含义如下：
    1.organ_code (varchar(255)): 组织代码，标识所属的公司或组织。
    2.organ_name (varchar(255)): 组织名称，标识公司或机构的名称。
    3.tdate (timestamp): 录入日期，表示该数据的记录时间。
    4.data_code (varchar(255)): 数据代码，标识该条记录的唯一代码。
    5.proj_name (varchar(255)): 科目名称，描述该记录所涉及的科目。
    6.item_id (float8): ID号，唯一标识符。
    7.unit (varchar(255)): 单位，表示数值单位。
    8.curr_y_budget (float8): 本年预算总金额，表示公司当前年度的预算总额。
    9.last_m_pe_num (float8): 上月完成总金额，表示上个月实际完成的金额。
    10.curr_m_pe_num (float8): 本月完成总金额，表示当前月份实际完成的金额。
    11.last_ym_pe_num (float8): 上年同期总金额，表示上一年相同时间段的实际完成金额。
    12.curr_y_pe_num (float8): 本年累计总金额，表示当前年度至今的累计完成金额。
    13.curr_y_same_pe_num (float8): 同比增减额总金额，表示同比增加或减少的金额。
    14.curr_y_same_pe_factor (float8): 同比增减因素影响原因，分析同比增减的原因。
    15.curr_y_budegt_same_num (float8): 本年预算同比增减单位金额，表示本年预算相较于同期的增减变化。
    16.curr_y_budegt_same_factor (float8): 本年预算同比增减的原因，分析预算变化的原因。
    17.uploadtime (timestamp): 上传日期，表示数据的最后更新时间。
    18.curr_y_budget_unit (float8): 本年预算单位金额，表示单位的预算金额。
    19.last_m_pe_num_unit (float8): 上月完成单位金额，表示上个月单位的实际完成金额。
    20.curr_m_pe_num_unit (float8): 本月完成单位金额，表示本月单位的实际完成金额。
    21.curr_y_pe_num_unit (float8): 本年累计单位金额，表示当前年度至今的单位累计完成金额。
    22.last_ym_pe_num_unit (float8): 上年同期单位金额，表示上一年相同时间段的单位完成金额。
    23.curr_y_same_pe_num_unit (float8): 同比增减额单位金额，表示单位同比增加或减少的金额。
    24.curr_y_same_pe_factor_unit (float8): 同比增减因素的影响金额，分析同比增减的具体数值。
    25.cause_factor (float8): 同比增减利因素的影响原因，分析同比增减变化的具体原因。\n"""
)
#
#
proj_name_7 = vn.run_sql("SELECT DISTINCT proj_name FROM twc_thermal_power_operate_m WHERE proj_name IS NOT NULL ")
organ_name_7 = vn.run_sql("SELECT DISTINCT organ_name FROM twc_thermal_power_operate_m WHERE organ_name IS NOT NULL ")

proj_name_7_list = [item for item in proj_name_7['proj_name'].tolist() if item is not None]
organ_name_7_list = [item for item in organ_name_7['organ_name'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
proj_name_7_str = ', '.join(proj_name_7_list)
organ_name_7_str = ', '.join((organ_name_7_list))
data_rag['twc_thermal_power_operate_m'].append(
    """
    twc_thermal_power_operate_m表
    """ +
    f"""**关键字段**:
1. **proj_name** (科目名):
   - **固定取值列表**: 
     - [proj_name_7_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_name** (组织名称):
   - **固定取值列表**:
     - [{organ_name_7_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。

**注意**:
- **proj_name** 和 **organ_name** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=proj_name_7_list, name='twc_thermal_power_operate_m')

# #twc_ygl0060
data_rag['twc_ygl0060'] = []
data_rag['twc_ygl0060'].append(
    """
    -- public.twc_ygl0060 definition
    
    -- Drop table
    
    -- DROP TABLE public.twc_ygl0060;
    
    CREATE TABLE public.twc_ygl0060 (
        organ_code varchar(255) NULL, -- 组织代码
        tdate timestamp NULL, -- 录入日期
        proj_name varchar(255) NULL, -- 科目名称
        item_id float8 NULL, -- id号
        unit varchar(255) NULL, -- 单位
        curr_y_budget float8 NULL, -- 本年预算总金额
        last_m_pe_num float8 NULL, -- 上月完成总金额
        curr_m_pe_num float8 NULL, -- 本月完成总金额
        last_ym_pe_num float8 NULL, -- 上年同期总金额
        curr_y_pe_num float8 NULL, -- 本年累计总金额
        curr_y_same_pe_num float8 NULL, -- 同比增减额总金额
        curr_y_same_pe_factor float8 NULL, -- 同比增减利因素的影响原因
        curr_y_budegt_same_num float8 NULL, -- 本年预算同比增减总金额
        curr_y_budegt_same_factor float8 NULL, -- 本年预算同比增减利因素的影响原因
        uploadtime timestamp NULL, -- 更新日期
        organ_name varchar(255) NULL, -- 组织名称
        data_code varchar(255) NULL, -- 数据代码
        curr_y_budget_unit float8 NULL, -- 本年预算单位总金额
        last_m_pe_num_unit float8 NULL, -- 上月完成单位金额
        curr_m_pe_num_unit float8 NULL, -- 本月完成单位金额
        curr_y_pe_num_unit float8 NULL, -- 本年累计单位金额
        last_ym_pe_num_unit float8 NULL, -- 上年同期单位金额
        curr_y_same_pe_num_unit float8 NULL, -- 同比增减额单位金额
        curr_y_same_pe_factor_unit float8 NULL, -- 同比增减利因素的影响金额
        cause_factor float8 NULL -- 同比增减利因素的影响原因
    );
    COMMENT ON TABLE public.twc_ygl0060 IS '新能源发电经营情况表';
    
    -- Column comments
    
    COMMENT ON COLUMN public.twc_ygl0060.organ_code IS '组织代码';
    COMMENT ON COLUMN public.twc_ygl0060.tdate IS '录入日期';
    COMMENT ON COLUMN public.twc_ygl0060.proj_name IS '科目名称';
    COMMENT ON COLUMN public.twc_ygl0060.item_id IS 'id号';
    COMMENT ON COLUMN public.twc_ygl0060.unit IS '单位';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_budget IS '本年预算总金额';
    COMMENT ON COLUMN public.twc_ygl0060.last_m_pe_num IS '上月完成总金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_m_pe_num IS '本月完成总金额';
    COMMENT ON COLUMN public.twc_ygl0060.last_ym_pe_num IS '上年同期总金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_pe_num IS '本年累计总金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_same_pe_num IS '同比增减额总金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_same_pe_factor IS '同比增减利因素的影响原因';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_budegt_same_num IS '本年预算同比增减总金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_budegt_same_factor IS '本年预算同比增减利因素的影响原因';
    COMMENT ON COLUMN public.twc_ygl0060.uploadtime IS '更新日期';
    COMMENT ON COLUMN public.twc_ygl0060.organ_name IS '组织名称';
    COMMENT ON COLUMN public.twc_ygl0060.data_code IS '数据代码';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_budget_unit IS '本年预算单位总金额';
    COMMENT ON COLUMN public.twc_ygl0060.last_m_pe_num_unit IS '上月完成单位金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_m_pe_num_unit IS '本月完成单位金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_pe_num_unit IS '本年累计单位金额';
    COMMENT ON COLUMN public.twc_ygl0060.last_ym_pe_num_unit IS '上年同期单位金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_same_pe_num_unit IS '同比增减额单位金额';
    COMMENT ON COLUMN public.twc_ygl0060.curr_y_same_pe_factor_unit IS '同比增减利因素的影响金额';
    COMMENT ON COLUMN public.twc_ygl0060.cause_factor IS '同比增减利因素的影响原因';
    
    """
)
data_rag['twc_ygl0060'].append(
    """twc_ygl0060是一张新能源发电经营情况表，
    该表用于记录新能源发电公司各项经营指标的预算、完成情况以及同比变化。
    通过该表，用户可以分析公司经营的进展、预算的执行情况，并追踪同比增减的因素及其影响。
    该表的主要列及其含义如下：
    1.organ_code (varchar(255)): 组织代码，标识所属的公司或组织。
    2.tdate (timestamp): 录入日期，表示数据的录入时间。
    3.proj_name (varchar(255)): 科目名称，描述该记录所涉及的科目。
    4.item_id (float8): ID号，唯一标识符。
    5.unit (varchar(255)): 单位，表示数值单位。
    6.curr_y_budget (float8): 本年预算总金额，表示公司当前年度的预算总额。
    7.last_m_pe_num (float8): 上月完成总金额，表示上个月实际完成的金额。
    8.curr_m_pe_num (float8): 本月完成总金额，表示当前月份实际完成的金额。
    9.last_ym_pe_num (float8): 上年同期总金额，表示上一年相同时间段的实际完成金额。
    10.curr_y_pe_num (float8): 本年累计总金额，表示当前年度至今的累计完成金额。
    11.curr_y_same_pe_num (float8): 同比增减额总金额，表示同比增加或减少的金额。
    12.curr_y_same_pe_factor (float8): 同比增减因素影响原因，分析同比增减的原因。
    13.curr_y_budegt_same_num (float8): 本年预算同比增减总金额，表示本年预算相较于同期的增减变化。
    14.curr_y_budegt_same_factor (float8): 本年预算同比增减利因素的影响原因，分析预算变化的原因。
    15.uploadtime (timestamp): 更新日期，表示数据的最后更新时间。
    16.organ_name (varchar(255)): 组织名称，标识公司或机构的名称。
    17.data_code (varchar(255)): 数据代码，标识该条记录的唯一代码。
    18.curr_y_budget_unit (float8): 本年预算单位总金额，表示单位的预算金额。
    19.last_m_pe_num_unit (float8): 上月完成单位金额，表示上个月单位的实际完成金额。
    20.curr_m_pe_num_unit (float8): 本月完成单位金额，表示本月单位的实际完成金额。
    21.curr_y_pe_num_unit (float8): 本年累计单位金额，表示当前年度至今的单位累计完成金额。
    22.last_ym_pe_num_unit (float8): 上年同期单位金额，表示上一年相同时间段的单位完成金额。
    23.curr_y_same_pe_num_unit (float8): 同比增减额单位金额，表示单位同比增加或减少的金额。
    24.curr_y_same_pe_factor_unit (float8): 同比增减因素的影响金额，分析同比增减的具体数值。
    25.cause_factor (float8): 同比增减利因素的影响原因，分析同比增减变化的具体原因\n"""
)

proj_name_8 = vn.run_sql("SELECT DISTINCT proj_name FROM twc_ygl0060 WHERE proj_name IS NOT NULL ")
organ_code_8 = vn.run_sql("SELECT DISTINCT organ_code FROM twc_ygl0060 WHERE organ_code IS NOT NULL ")

proj_name_8_list = [item for item in proj_name_8['proj_name'].tolist() if item is not None]
organ_code_8_list = [item for item in organ_code_8['organ_code'].tolist() if item is not None]
# print(item_names_list)
# 将列表中的内容转换为字符串，用逗号分隔
proj_name_8_str = ', '.join(proj_name_8_list)
organ_code_8_str = ', '.join((organ_code_8_list))

data_rag['twc_ygl0060'].append(
    """
    twc_ygl0060表
    """ +
    f"""**关键字段**:
1. **proj_name** (科目名):
   - **固定取值列表**: 
     - [proj_name_8_str]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
2. **organ_code** (组织名称):
   - **固定取值列表**:
     - [{organ_code_8_str}]
   - **说明**: 该字段的值只能是上述列表中的某一个，且不能有其他取值。
**注意**:
- **proj_name** 和 **organ_code** 必须严格按照上述固定取值中的一个进行匹配。
- 任何不在上述固定列表中的值都是无效的，不应出现在查询的 `WHERE` 子句中。
"""
)
vn.train_yu(text=proj_name_8_list, name='twc_ygl0060')
tmp_question = "，请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL，如果不需要，请直接回答，如果需要生成sql，生成一次sql，请严格按照system内容的固定取值列表来生成sql，注意一定要在sql中给出结果。"
# endregion

import os.path
import sqlite3
import time
from functools import lru_cache
from typing import List

from gradio_client import Client
from difflib import SequenceMatcher

db_path = "/code/kotaemon/ktem_app_data/user_data/sql.db"


# region 这个是用于实现知识库问答的函数，目前是直接通过api的方式，对kotaemon的功能进行调用，实际上应该把kotaemon的rag代码提取出来，放到这里
@lru_cache(1)  # 这个是自定义函数
def get_all_file_ids(db_path: str, indices: List[int] | None = None):
    assert os.path.isfile(db_path)

    file_ids: List[str] = []

    conn = sqlite3.connect(db_path)

    try:
        c = conn.cursor()
        if indices is None:
            # Get all fileIndex
            c.execute("SELECT * FROM ktem__index")
            indices = [each[0] for each in c.fetchall()]

        # For each index get all file_ids
        for i in indices:
            table_name = f"index__{i}__source"
            c.execute(
                f"SELECT * FROM {table_name}",
            )
            file_ids += [each[0] for each in c.fetchall()]
    finally:
        conn.close()

    return file_ids


def find_differences(str1, str2):
    matcher = SequenceMatcher(None, str1, str2)
    differences = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != 'equal':
            differences.append(str2[j1:j2])
    return "".join(differences)


def chat_rag(question):
    file_ids = get_all_file_ids(db_path)

    print(f"File_ids: {file_ids}")

    client = Client("http://localhost:7860/")

    job = client.submit(
        [
            [
                question + ",用中问回答",
                None,
            ]
        ],
        "select",
        file_ids,
        api_name="/chat_fn",
    )

    resps = ['']
    full_response = ""

    # 逐步收集响应并保存到文件
    for o in job:
        # 检查是否有新的内容返回
        if 'Thinking' not in o[0][0][1]:
            if len(resps[-1]) != len(o[0][0][1]):
                content = find_differences(resps[-1], o[0][0][1])
                resps.append(o[0][0][1])
                full_response += content

            else:
                break

    return full_response


# endregion
# region 这个是用于生成ppt的函数，是借鉴于slide_deck项目里面的ppt生成方法
def generate_slide_deck(json_str: str):
    """
    Create a slide deck.

    :param json_str: The content in *valid* JSON format.
    """

    path = "financial_report.pptx"
    pptx_template = r"./pptx_templates/Blank.pptx"

    try:
        ppt = pptx_helper.generate_powerpoint_presentation(
            json_str,
            slides_template=pptx_template,
            output_file_path=path
        )  # 生成幻灯片文件
    except ValueError:

        ppt = pptx_helper.generate_powerpoint_presentation(
            text_helper.fix_malformed_json(json_str),
            slides_template=pptx_template,
            output_file_path=path
        )  # 生成幻灯片文件
    return ppt


# endregion
def requires_cache(fields):  # 这个函数和缓存进行搭配
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            id = request.args.get('id')

            if id is None:
                return jsonify({"type": "error", "error": "No id provided"})

            for field in fields:
                if cache.get(id=id, field=field) is None:
                    return jsonify({"type": "error", "error": f"No {field} found"})

            field_values = {field: cache.get(id=id, field=field) for field in fields}

            # Add the id to the field_values
            field_values['id'] = id

            return f(*args, **field_values, **kwargs)

        return decorated

    return decorator


# 下面是一些函数的定义，主要是和前端进行交互，下面一一进行介绍
# region 这个函数是用于界面刚刷新的时候，生成一些问题，目前点击问题是没有任何效果的
@app.route('/api/v0/generate_questions', methods=['GET'])
def generate_questions():
    # vn.generate_questions()
    return jsonify({
        "type": "question_list",
        "questions": ["告诉我H公司在23年4月份，营业收入的本月完成总金额是多少", "告诉我D公司合并在22年11月份，本月的综合收益总额是多少"],
        "header": "这里有一些您可能想查询的数据:"
    })


# endregion
# region 这个函数是和前端交互，生成sql答案的，这里面也包含了sql自纠正
@app.route('/api/v0/generate_sql', methods=['GET'])
def generate_sql():
    question = flask.request.args.get('question')
    former_doc_list = flask.request.args.get("former_doc_list")
    mode = flask.request.args.get("mode_web")
    id = cache.generate_id(question=question)

    global doc_list_tmp
    global ddl_list_tmp
    # question=question+"。你是一个数据库和知识系统的助手，你可以根据数据库中的表来回答数据相关问题，也可以通过你的知识回答一般性问题。请在回答问题前，判断该问题是否需要查询数据库。如果需要查询数据库，请生成SQL,确认好表的字段，不要搞错了；如果不需要，请直接回答,在生成sql时将WHERE子句中使用的所有字段添加到 SELECT中，生成的sql，用like代替等于号，同时like要搭配%符号，日期字段的属性是timestamp"
    if question is None:
        return jsonify({"type": "error", "error": "No question provided"})

    if mode == "file":
        return jsonify(
            {
                "type": "sql",
                "id": id,
                "sql_que": "",
                "text": chat_rag(question),  # 文档查询的
            })

    else:
        if former_doc_list == "true":

            response = vn.generate_sql(question=question, allow_llm_to_see_data=True, history=history,
                                       doc_list=doc_list_tmp, ddl_list=ddl_list_tmp)
        else:
            # question+="，根据上下文判断是否需要生成sql，如果需要生成sql，就只生成一次sql，在sql中得出答案"
            response = vn.generate_sql(question=question, allow_llm_to_see_data=True, history=history, doc_list=[])
            doc_list_tmp = response[2]
            ddl_list_tmp = response[3]
        if response[0] != "":
            cache.set(id=id, field='question', value=question)
            cache.set(id=id, field='sql', value=response[0])
        history.extend([{"role": "user", "content": question}, {"role": "assistant", "content": response[0]}])

        return jsonify(
            {
                "type": "sql",
                "id": id,
                "sql_que": response[0],
                "text": response[1],
            })


# endregion
# region 这个函数是用于和前端交互，生成word文档的
@app.route('/api/v0/generate_word', methods=['POST'])
def generate_word():
    data = flask.request.get_json()

    # 获取 reportList，确保后端正确处理传来的数据
    history_for_word = data.get('reportList')
    message_log = vn.generate_message_log(history_for_word)
    print(message_log)
    word_text = vn.submit_prompt(message_log)
    word_text = vn.create_word_from_markdown(word_text)

    output = BytesIO()
    # output.write(word_text.encode('utf-8'))  # 转为字节流
    word_text.save(output)
    output.seek(0)  # 重置流的位置，以确保从头读取

    # 返回 Word 文档作为文件
    return Response(
        output.getvalue(),  # 返回字节数据
        mimetype="application/msword",  # 或 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        headers={"Content-Disposition": "attachment; filename=financial_report.docx"}
    )


# endregion
# region 这个函数和前端交互，用于生成ppt
@app.route('/api/v0/generate_PPT', methods=['POST'])
def generate_PPT():
    data = flask.request.get_json()

    # 获取 reportList，确保后端正确处理传来的数据
    history_for_word = data.get('reportList')
    message_log = vn.generate_message_log_ppt(history_for_word)
    print(message_log)
    response = vn.submit_prompt(message_log)
    response_cleaned = text_helper.get_clean_json(response)
    generate_ppt = generate_slide_deck(response_cleaned)
    # presentation = pptx.Presentation
    # 将生成的 pptx 文件保存在内存中
    ppt_io = BytesIO()
    generate_ppt.save(ppt_io)
    ppt_io.seek(0)  # 将指针移动到文件开头

    return Response(
        ppt_io.getvalue(),  # 返回字节数据
        mimetype="application/vnd.openxmlformats-officedocument.presentationml.presentation",  # MIME 类型
        headers={"Content-Disposition": "attachment; filename=generated_presentation.pptx"}  # 设置下载文件名
    )


# endregion

# region 这函数是用于执行sql语句
@app.route('/api/v0/run_sql', methods=['GET'])
@requires_cache(['sql'])
def run_sql(id: os, sql: os):
    # def run_sql():
    try:
        # id = flask.request.args.get('id')
        # sql = flask.request.args.get('sql')
        df = vn.run_sql(sql=sql)
        for col in df.columns:
            if pd.api.types.is_datetime64_any_dtype(df[col]):
                df[col] = df[col].dt.strftime('%Y-%m-%d')  # 修改日期格式为 "YYYY-MM-DD"
            elif pd.api.types.is_object_dtype(df[col]):
                try:
                    temp_col = pd.to_datetime(df[col], errors='coerce')  # 将字符串列转为日期，无法转的变为 NaT
                    # 如果转换后有有效的日期，则格式化，否则保持原样
                    if temp_col.notna().any():
                        df[col] = temp_col.dt.strftime('%Y-%m-%d')  # 格式化有效日期
                except Exception as e:
                    print(f"Error in converting column {col}: {e}")

        cache.set(id=id, field='df', value=df)
    except Exception as e:
        error_message = "SQL执行错误,重新生成:" + str(e)
        print(error_message)
        return jsonify({"type": "error", "error": error_message})
    # history[-1]["content"]+="\n\n查询结果如下：\n\n"+df.to_markdown(index=False)
    # history.extend([{"role": "user", "content":"\n\n查询结果如下：\n\n"+df.to_markdown(index=False)},
    #                 {"role": "assistant", "content": "好的"}])
    return jsonify(
        {
            "type": "df",
            "id": id,
            "df": df.to_json(orient='records'),
        })

    # except Exception as e:
    #     return jsonify({"type": "error", "error": os(e)})


# endregion
# region 这个函数是用于前端进行csv的下载
@app.route('/api/v0/download_csv', methods=['GET'])
@requires_cache(['df'])
def download_csv(id: os, df):
    output = BytesIO()
    # 将 CSV 编码为 GBK 并写入字节流
    df.to_csv(output, encoding="GBK", index=False)
    output.seek(0)  # 重置流的位置，以确保从头读取

    return Response(
        output.getvalue(),  # 返回字节数据
        mimetype="text/csv",
        headers={"Content-Disposition": f"attachment; filename={id}.csv"}
    )


# endregion
# region 这个函数是用于前端图表的生成
@app.route('/api/v0/generate_plotly_figure', methods=['GET'])
@requires_cache(['df', 'sql'])
def generate_plotly_figure(id: os, df, sql):
    question = flask.request.args.get('question')
    try:

        code = vn.generate_plotly_code(question=question, sql=sql,
                                       df_metadata=f"Running df.dtypes gives:\n {df.dtypes}")
        print(code)
        summery = vn.generate_summary(question, df)
        fig = vn.get_plotly_figure(plotly_code=code, df=df, dark_mode=False)
        fig_json = fig.to_json()

        cache.set(id=id, field='fig_json', value=fig_json)

        return jsonify(
            {
                "type": "plotly_figure",
                "id": id,
                "fig": fig_json,
                "summary": summery
            })
    except Exception as e:
        # Print the stack trace
        import traceback
        traceback.print_exc()

        return jsonify({"type": "error", "error": os(e)})


# endregion
# region 这个代码是用于获取训练数据的函数，这个是Vanna原本的函数，不过现在已经不需要这个函数了
@app.route('/api/v0/get_training_data', methods=['GET'])
def get_training_data():
    df = vn.get_training_data()

    return jsonify(
        {
            "type": "df",
            "id": "training_data",
            "df": df.head(25).to_json(orient='records'),
        })


# endregion
# region 这个函数是移除训练的数据，也已经不需要了
@app.route('/api/v0/remove_training_data', methods=['POST'])
def remove_training_data():
    # Get id from the JSON body
    id = flask.request.json.get('id')

    if id is None:
        return jsonify({"type": "error", "error": "No id provided"})

    if vn.remove_training_data(id=id):
        return jsonify({"success": True})
    else:
        return jsonify({"type": "error", "error": "Couldn't remove training data"})


# endregion
# region 这个函数是用于将数据进行训练，但这个已经不需要使用了
@app.route('/api/v0/train', methods=['POST'])
def add_training_data():
    question = flask.request.json.get('question')
    sql = flask.request.json.get('sql')
    ddl = flask.request.json.get('ddl')
    documentation = flask.request.json.get('documentation')

    try:
        id = vn.train(question=question, sql=sql, ddl=ddl, documentation=documentation)

        return jsonify({"id": id})
    except Exception as e:
        print("TRAINING ERROR", e)
        return jsonify({"type": "error", "error": os(e)})


# endregion
# region 同样这个也不需要使用了
@app.route('/api/v0/generate_followup_questions', methods=['GET'])
@requires_cache(['df', 'question', 'sql'])
def generate_followup_questions(id: os, df, question, sql):
    followup_questions = vn.generate_followup_questions(question=question, sql=sql, df=df)

    cache.set(id=id, field='followup_questions', value=followup_questions)

    return jsonify(
        {
            "type": "question_list",
            "id": id,
            "questions": followup_questions,
            "header": "Here are some followup questions you can ask:"
        })


# endregion
# region同样也不需要这个函数
@app.route('/api/v0/load_question', methods=['GET'])
@requires_cache(['question', 'sql', 'df', 'fig_json', 'followup_questions'])
def load_question(id: os, question, sql, df, fig_json, followup_questions):
    try:
        return jsonify(
            {
                "type": "question_cache",
                "id": id,
                "question": question,
                "sql": sql,
                "df": df.head(10).to_json(orient='records'),
                "fig": fig_json,
                "followup_questions": followup_questions,
            })

    except Exception as e:
        return jsonify({"type": "error", "error": os(e)})


# endregion
# region这个是用于在界面显示历史会话，但现在暂时没使用这个函数
@app.route('/api/v0/get_question_history', methods=['GET'])
def get_question_history():
    return jsonify({"type": "question_history", "questions": cache.get_all(field_list=['question'])})


# endregion
@app.route('/')
def root():
    return app.send_static_file('index.html')


if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=False)
    # print(vn.submit_prompt(model="qwq:32b",stream=False,prompt=[{"role": "user","content": "你知道sql如何写吗，假如你是一个财务总监"}]))
    # npm create vite@latest my-svelte-app
    # npm install svelte@5.0.0
