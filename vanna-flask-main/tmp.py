from docx import Document


def generate_financial_report(query, results, analysis, suggestions):
    # 创建一个新的 Word 文档
    doc = Document()

    # 添加前言
    doc.add_heading('前言', level=1)
    doc.add_paragraph("本报告基于用户的查询，旨在分析财务指标，为后续决策提供支持。")

    # 添加查询内容
    doc.add_heading('查询内容', level=1)
    doc.add_paragraph(f"查询内容：{query}")

    # 添加数据展示
    doc.add_heading('数据展示', level=1)
    for item in results:
        doc.add_paragraph(f"{item['时间']} - {item['指标']}: {item['值']} {item['单位']}")

    # 添加深入分析
    doc.add_heading('深入分析', level=1)
    doc.add_paragraph(analysis)

    # 添加结论与建议
    doc.add_heading('结论与建议', level=1)
    doc.add_paragraph(suggestions)

    # 保存文档
    file_path = "financial_report.docx"
    doc.save(file_path)
    return file_path


def create_word_from_markdown(markdown_text):
    """
    将 Markdown 文本内容解析为 Word 文档。
    :param markdown_text: 大模型生成的 Markdown 文本
    :return: Word 文档对象
    """
    document = Document()
    markdown_text = markdown_text.replace("markdown", "")
    markdown_text = markdown_text.replace("```", "")
    lines = markdown_text.split("\n")
    print(lines)
    for line in lines:
        if line.startswith("# "):  # 一级标题
            document.add_heading(line[2:], level=1)
        elif line.startswith("### "):  # 二级标题
            document.add_heading(line[4:], level=1)
        elif line.startswith("- "):  # 列表项
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():  # 普通段落
            document.add_paragraph(line.strip())
    file_path = "financial_report.docx"
    document.save(file_path)
    return file_path


query = "清远电厂的资产负债率"
results = [
    {"时间": "2023年1月", "指标": "资产负债率", "值": 0.65, "单位": ""},
    {"时间": "2023年2月", "指标": "资产负债率", "值": 0.66, "单位": ""},
]
text = """
```markdown
### 前言
查询背景和目的：为了了解A公司在2023年的经营状况，特别是在合并后的财务表现方面，需要分析其每个月的利润总额变化情况。通过此报告可以识别出潜在的问题区域，并为进一步的战略规划提供数据支持。
### 查询内容
用户请求的查询项为：
•	A公司合并在2023年每个月的利润总额
### 数据展示
| 月份       | 利润总额（万元） |
|------------|------------------|
| 2023-01    | 1648.16          |
| 2023-02    | 12618.90         |
| 2023-03    | 22615.20         |
| 2023-04    | 18995.56         |
| 2023-05    | 24913.40         |
| 2023-06    | 19866.12         |
| 2023-07    | 28204.21         |
### 深入分析
1. **利润总额变化趋势**：
- 利润从一月份开始上升，二月份显著增长。这表明公司在前两个月期间可能有重大收入增加或成本减少。
- 在三月和四月间, 利润下降但仍然保持在较高水平。
- 五月利润再次上涨, 表明公司运营情况良好且收益稳定。
2. **季节性因素**：
- 某些月份的利润大幅波动可能与特定产品线或服务的周期性需求有关，如夏季旅游高峰期间的服务业收入增加等。
3. **外部环境影响**：
- 由于缺乏具体财务报表，很难确定是否有经济政策调整、市场竞争加剧等因素对A公司的利润产生显著影响。需要进一步分析其他相关财务指标（如资产和负债）以获得更全面的了解。
4. **财务健康度评估**：
- 资产：通过资产总额的变化来反映公司扩张情况，可能与收入增长有关联；
- 负债：观察到利润上升同时需留意债务水平是否合理，过高时将会影响净利润率和偿债能力；
- 所有者权益变化：反映了公司的资本积累状况。
### 结论与建议
•	总体来看，A公司在2023年的合并后表现出稳定的盈利增长趋势。
•	为确保未来可持续发展, 建议公司持续监控外部经济环境和内部运营效率。同时应关注潜在风险因素如高负债率对公司财务稳定性的影响。
请根据上述报告进行进一步分析并制定详细的行动计划。
```
"""
analysis = "资产负债率在2023年初略有上升，主要原因是负债增加较快。"
suggestions = "建议加强资金管理，优化资产配置。"

# file_path = generate_financial_report(query, results, analysis, suggestions)
file_path = create_word_from_markdown(text)
print(f"文档已保存到：{file_path}")
